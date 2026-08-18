# -*- coding: utf-8 -*-
"""
Word 检测报告生成模块（纯本地 python-docx）
=========================================
报告结构：
    封面 → 检测基本信息 → 整体统计汇总 → 问题类型分布
    → 分文件错误明细 → 异常（无法解析）文件清单 → 整改建议 → 落款

排版特性：
    A4 页面、页边距规范、中文字体（微软雅黑 + 宋体）
    带边框表格、斑马纹表头底色、页码页脚
    可直接归档打印

保密说明：仅本地生成 docx 文件，不联网、不上传、不写入任何外部位置。
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from checkers.base import STATUS_UNREADABLE, FileResult
from config.config_manager import severity_label

# ---------------------------------------------------------------------------
# 样式常量
# ---------------------------------------------------------------------------
FONT_TITLE = "微软雅黑"
FONT_BODY = "宋体"
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)     # 沉稳科技蓝
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x80, 0x80, 0x80)
COLOR_HIGH = RGBColor(0xC0, 0x39, 0x2B)
COLOR_MEDIUM = RGBColor(0xB9, 0x7A, 0x0A)
COLOR_LOW = RGBColor(0x55, 0x7A, 0x95)

SHADE_HEADER = "DCE6F1"    # 表头底色（浅蓝）
SHADE_SUBTLE = "F2F5F9"    # 次级底色

SEVERITY_COLOR = {"high": COLOR_HIGH, "medium": COLOR_MEDIUM, "low": COLOR_LOW}


# ---------------------------------------------------------------------------
# 底层排版辅助函数
# ---------------------------------------------------------------------------
def _set_run_font(run: Any, name: str = FONT_BODY, size: float = 10.5,
                  bold: bool = False, color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（含东亚字体，保证中文正确显示）。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 东亚字体必须写入 rFonts/eastAsia，否则中文可能回退为默认字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _para(doc_or_cell: Any, text: str = "", size: float = 10.5, bold: bool = False,
          align: Optional[int] = None, font: str = FONT_BODY,
          color: Optional[RGBColor] = None, space_after: float = 6,
          space_before: float = 0, line_spacing: Optional[float] = None,
          first_line_indent: Optional[float] = None) -> Any:
    """统一的段落创建函数。"""
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if line_spacing:
        pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold, color)
    return p


def _shade_cell(cell: Any, hex_color: str) -> None:
    """给表格单元格设置底纹颜色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell: Any, text: str, size: float = 9.5, bold: bool = False,
                   align: Optional[int] = None, color: Optional[RGBColor] = None,
                   font: str = FONT_BODY) -> None:
    """写入单元格文本（清空原有内容）。"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("" if text is None else str(text))
    _set_run_font(run, font, size, bold, color)


def _set_repeat_header(row: Any) -> None:
    """设置表格标题行跨页重复（报告自身也要符合规范）。"""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _set_col_widths(table: Any, widths_cm: List[float]) -> None:
    """固定列宽（需对每个单元格设置才生效）。"""
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def _add_page_number_footer(section: Any) -> None:
    """页脚添加「第 X 页 共 Y 页」域代码。"""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr: str) -> None:
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = instr
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_el)
        run._r.append(fld_end)
        _set_run_font(run, FONT_BODY, 9, color=COLOR_MUTED)

    r1 = p.add_run("第 ")
    _set_run_font(r1, FONT_BODY, 9, color=COLOR_MUTED)
    _field("PAGE")
    r2 = p.add_run(" 页  共 ")
    _set_run_font(r2, FONT_BODY, 9, color=COLOR_MUTED)
    _field("NUMPAGES")
    r3 = p.add_run(" 页")
    _set_run_font(r3, FONT_BODY, 9, color=COLOR_MUTED)


def _h1(doc: Any, text: str) -> None:
    """一级标题：蓝色粗体 + 下方分隔横线效果。"""
    p = _para(doc, text, size=15, bold=True, font=FONT_TITLE,
              color=COLOR_PRIMARY, space_before=14, space_after=8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc: Any, text: str) -> None:
    """二级标题。"""
    _para(doc, text, size=12, bold=True, font=FONT_TITLE,
          color=COLOR_TEXT, space_before=10, space_after=5)


def _kv_table(doc: Any, rows: List[tuple], key_width: float = 4.5,
              val_width: float = 11.5) -> Any:
    """键值信息表。"""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for key, val in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], key, size=10, bold=True)
        _shade_cell(cells[0], SHADE_SUBTLE)
        _set_cell_text(cells[1], val, size=10)
    _set_col_widths(table, [key_width, val_width])
    return table


# ---------------------------------------------------------------------------
# 报告主体
# ---------------------------------------------------------------------------
class ReportBuilder:
    """检测报告生成器。"""

    def __init__(self, results: List[FileResult], summary: Dict[str, Any],
                 rule_summary: Optional[Dict[str, Any]] = None,
                 operator: str = "", org: str = "",
                 include_cover: bool = True,
                 detail_columns: Optional[List[str]] = None) -> None:
        self.results = results
        self.summary = summary
        self.rule_summary = rule_summary or {}
        self.operator = operator or "—"
        self.org = org or "—"
        self.include_cover = include_cover
        self.detail_columns = detail_columns or ["index", "location", "type", "severity", "detail", "suggestion"]
        self.now = datetime.now()

    # ------------------------------------------------------------------
    def build(self, out_path: str) -> str:
        """生成报告并保存到本地路径，返回最终文件路径。"""
        doc = Document()
        self._setup_page(doc)
        if self.include_cover:
            self._cover(doc)
        self._basic_info(doc)
        self._overall_stats(doc)
        self._rule_distribution(doc)
        self._file_details(doc)
        self._unreadable_list(doc)
        self._suggestions(doc)
        self._footer_note(doc)

        out_dir = os.path.dirname(os.path.abspath(out_path))
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        doc.save(out_path)
        return out_path

    # ------------------------------------------------------------------
    def _setup_page(self, doc: Any) -> None:
        """A4 页面 + 页边距 + 默认样式 + 页脚页码。"""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        _add_page_number_footer(section)

        style = doc.styles["Normal"]
        style.font.name = FONT_BODY
        style.font.size = Pt(10.5)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # ------------------------------------------------------------------
    def _cover(self, doc: Any) -> None:
        """报告封面。"""
        _para(doc, "", space_after=60)
        _para(doc, "文档低级错误检查报告", size=30, bold=True, font=FONT_TITLE,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_PRIMARY, space_after=10)
        _para(doc, "Document Quality Inspection Report", size=12, font=FONT_TITLE,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED, space_after=8)
        _para(doc, "— 本地离线检测 · 保密模式 —", size=11, font=FONT_TITLE,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED, space_after=50)

        total = int(self.summary.get("total_files", 0))
        issues = int(self.summary.get("total_issues", 0))
        issue_files = int(self.summary.get("issue_files", 0))
        err_files = int(self.summary.get("error_files", 0))

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cover_rows = [
            ("检测文件总数", f"{total} 个"),
            ("存在问题文件", f"{issue_files} 个"),
            ("无法解析文件", f"{err_files} 个"),
            ("发现问题总数", f"{issues} 处"),
            ("检测时间", self.now.strftime("%Y 年 %m 月 %d 日 %H:%M")),
            ("检测人", self.operator),
            ("所属单位", self.org),
        ]
        for key, val in cover_rows:
            cells = table.add_row().cells
            _set_cell_text(cells[0], key, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade_cell(cells[0], SHADE_HEADER)
            _set_cell_text(cells[1], val, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_col_widths(table, [5.5, 7.5])

        _para(doc, "", space_after=40)
        _para(doc, "本报告由本地离线检测工具自动生成，全过程无网络连接、无数据外发，",
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED, space_after=2)
        _para(doc, "所有文档内容仅在本机内存与本地磁盘处理，报告仅供内部核查使用。",
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED, space_after=0)

        # 封面结束分页
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------------------
    def _basic_info(self, doc: Any) -> None:
        """一、检测基本信息。"""
        _h1(doc, "一、检测基本信息")
        word_cnt = int(self.summary.get("by_type", {}).get("Word", 0))
        excel_cnt = int(self.summary.get("by_type", {}).get("Excel", 0))
        pdf_cnt = int(self.summary.get("by_type", {}).get("PDF", 0))
        other_cnt = int(self.summary.get("by_type", {}).get("不支持", 0))

        rules_word = self.rule_summary.get("word", "—")
        rules_excel = self.rule_summary.get("excel", "—")
        rules_textnorm = self.rule_summary.get("textnorm", "—")
        rules_pdf = self.rule_summary.get("pdf", "—")

        scope_txt = f"Word 文档 {word_cnt} 个、Excel 文档 {excel_cnt} 个、PDF 文档 {pdf_cnt} 个"
        if other_cnt:
            scope_txt += f"、不支持格式 {other_cnt} 个"

        _kv_table(doc, [
            ("检测工具", "文档低级错误检查工具（本地离线版 v1.0）"),
            ("运行模式", "纯本地离线运行，零联网、零数据上传、零云端接口"),
            ("检测时间", self.now.strftime("%Y-%m-%d %H:%M:%S")),
            ("检测范围", scope_txt),
            ("启用规则", f"Word 规则 {rules_word}；Excel 规则 {rules_excel}；"
                        f"文字规范规则 {rules_textnorm}；PDF 规则 {rules_pdf}"),
            ("检测方式", "本地规则校验（不调用任何 AI 云端服务、不解析上传）"),
            ("原文档处理", "只读检测，不对原始文档做任何写入或修改"),
            ("检测人", self.operator),
            ("所属单位", self.org),
        ])
        _para(doc, "", space_after=4)

    # ------------------------------------------------------------------
    def _overall_stats(self, doc: Any) -> None:
        """二、整体统计汇总。"""
        _h1(doc, "二、整体统计汇总")

        total = int(self.summary.get("total_files", 0))
        ok = int(self.summary.get("pass_files", 0))
        bad = int(self.summary.get("issue_files", 0))
        err = int(self.summary.get("error_files", 0))
        issues = int(self.summary.get("total_issues", 0))
        sev = self.summary.get("severity", {}) or {}
        ignored = int(self.summary.get("ignored_count", 0))
        checked = int(self.summary.get("checked_count", 0))

        rate = f"{(ok / total * 100):.1f}%" if total else "—"

        _h2(doc, "2.1 文件维度统计")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["文件总数", "检测通过", "存在问题", "无法解析", "通过率"]
        for idx, head in enumerate(headers):
            _set_cell_text(table.rows[0].cells[idx], head, size=10, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_TITLE)
            _shade_cell(table.rows[0].cells[idx], SHADE_HEADER)
        _set_repeat_header(table.rows[0])
        cells = table.add_row().cells
        for idx, val in enumerate([f"{total}", f"{ok}", f"{bad}", f"{err}", rate]):
            _set_cell_text(cells[idx], val, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_col_widths(table, [3.0, 3.0, 3.0, 3.0, 3.0])

        _para(doc, "", space_after=6)
        _h2(doc, "2.2 问题维度统计")
        table2 = doc.add_table(rows=1, cols=5)
        table2.style = "Table Grid"
        headers2 = ["问题总数", "严重", "一般", "轻微", "已忽略 / 已核查"]
        for idx, head in enumerate(headers2):
            _set_cell_text(table2.rows[0].cells[idx], head, size=10, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_TITLE)
            _shade_cell(table2.rows[0].cells[idx], SHADE_HEADER)
        _set_repeat_header(table2.rows[0])
        vals = [f"{issues}", f"{int(sev.get('high', 0))}", f"{int(sev.get('medium', 0))}",
                f"{int(sev.get('low', 0))}", f"{ignored} / {checked}"]
        cells2 = table2.add_row().cells
        for idx, val in enumerate(vals):
            color = None
            if idx == 1 and int(sev.get("high", 0)) > 0:
                color = COLOR_HIGH
            _set_cell_text(cells2[idx], val, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=color)
        _set_col_widths(table2, [3.0, 3.0, 3.0, 3.0, 3.0])

        _para(doc, "", space_after=6)
        concl = (f"本次共检测 {total} 个文件，其中 {ok} 个未发现低级错误，{bad} 个存在待整改问题，"
                 f"{err} 个无法解析；累计发现问题 {issues} 处"
                 f"（严重 {int(sev.get('high', 0))} 处、一般 {int(sev.get('medium', 0))} 处、"
                 f"轻微 {int(sev.get('low', 0))} 处）。")
        _para(doc, concl, size=10.5, line_spacing=1.5, first_line_indent=21)

    # ------------------------------------------------------------------
    def _rule_distribution(self, doc: Any) -> None:
        """三、问题类型分布。"""
        _h1(doc, "三、问题类型分布")
        by_rule: Dict[str, int] = dict(self.summary.get("by_rule", {}) or {})
        if not by_rule:
            _para(doc, "本次检测未发现任何问题，无问题类型分布数据。", size=10.5,
                  line_spacing=1.5, first_line_indent=21)
            return

        total = sum(by_rule.values()) or 1
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["序号", "问题类型", "出现次数", "占比"]
        for idx, head in enumerate(headers):
            _set_cell_text(table.rows[0].cells[idx], head, size=10, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_TITLE)
            _shade_cell(table.rows[0].cells[idx], SHADE_HEADER)
        _set_repeat_header(table.rows[0])

        for i, (name, count) in enumerate(by_rule.items(), start=1):
            cells = table.add_row().cells
            _set_cell_text(cells[0], str(i), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(cells[1], name, size=9.5)
            _set_cell_text(cells[2], str(count), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(cells[3], f"{count / total * 100:.1f}%", size=9.5,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            if i % 2 == 0:
                for c in cells:
                    _shade_cell(c, SHADE_SUBTLE)
        _set_col_widths(table, [1.6, 7.4, 3.0, 3.0])

    # ------------------------------------------------------------------
    def _file_details(self, doc: Any) -> None:
        """四、分文件错误明细。"""
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _h1(doc, "四、分文件错误明细")

        detail_targets = [r for r in self.results
                          if r.status != STATUS_UNREADABLE and r.active_issue_count > 0]
        if not detail_targets:
            _para(doc, "所有可解析文件均未发现低级错误，无明细清单。", size=10.5,
                  line_spacing=1.5, first_line_indent=21)
            return

        for f_idx, res in enumerate(detail_targets, start=1):
            sev = res.severity_count()
            _h2(doc, f"4.{f_idx} {res.file_name}")
            _para(doc,
                  f"文件类型：{res.file_type}    文件大小：{res.size_text}    "
                  f"问题总数：{res.active_issue_count} 处"
                  f"（严重 {sev['high']} / 一般 {sev['medium']} / 轻微 {sev['low']}）",
                  size=9.5, color=COLOR_MUTED, space_after=4)
            _para(doc, f"文件路径：{res.file_path}", size=8.5, color=COLOR_MUTED, space_after=6)

            cols = self.detail_columns or ["index", "location", "type", "severity", "detail", "suggestion"]
            col_meta = {
                "index": "序号", "location": "位置", "type": "问题类型",
                "severity": "级别", "detail": "问题说明 / 原文", "suggestion": "整改建议",
            }
            col_widths = {
                "index": 1.1, "location": 2.4, "type": 2.4, "severity": 1.1,
                "detail": 4.5, "suggestion": 3.5,
            }
            headers = [col_meta.get(c, c) for c in cols]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for idx, head in enumerate(headers):
                _set_cell_text(table.rows[0].cells[idx], head, size=9, bold=True,
                               align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_TITLE)
                _shade_cell(table.rows[0].cells[idx], SHADE_HEADER)
            _set_repeat_header(table.rows[0])

            shown = [i for i in res.issues if i.state != "ignored"]
            for i_idx, issue in enumerate(shown, start=1):
                cells = table.add_row().cells
                for c_idx, col in enumerate(cols):
                    if col == "index":
                        _set_cell_text(cells[c_idx], str(i_idx), size=8.5,
                                       align=WD_ALIGN_PARAGRAPH.CENTER)
                    elif col == "location":
                        _set_cell_text(cells[c_idx], issue.location, size=8.5)
                    elif col == "type":
                        _set_cell_text(cells[c_idx], issue.rule_title, size=8.5)
                    elif col == "severity":
                        _set_cell_text(cells[c_idx], severity_label(issue.severity), size=8.5,
                                       align=WD_ALIGN_PARAGRAPH.CENTER,
                                       color=SEVERITY_COLOR.get(issue.severity, COLOR_TEXT))
                    elif col == "detail":
                        detail_text = issue.detail
                        if issue.snippet:
                            detail_text += f"\n原文：{issue.snippet}"
                        _set_cell_text(cells[c_idx], detail_text, size=8.5)
                    elif col == "suggestion":
                        _set_cell_text(cells[c_idx], issue.suggestion or "—", size=8.5)
                if issue.state == "checked":
                    for c in cells:
                        _shade_cell(c, "EAF3EA")   # 已核查行浅绿底
            _set_col_widths(table, [col_widths.get(c, 2.0) for c in cols])

            if res.truncated:
                _para(doc, "注：该文件问题数量已达单文件上限，仅列出前述部分，建议整改后重新检测。",
                      size=8.5, color=COLOR_HIGH, space_before=3, space_after=8)
            else:
                _para(doc, "", space_after=8)

    # ------------------------------------------------------------------
    def _unreadable_list(self, doc: Any) -> None:
        """五、异常文件清单。"""
        _h1(doc, "五、异常文件清单（无法解析）")
        bad = [r for r in self.results if r.status == STATUS_UNREADABLE]
        if not bad:
            _para(doc, "本次检测所有文件均成功解析，无异常文件。", size=10.5,
                  line_spacing=1.5, first_line_indent=21)
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["序号", "文件名", "类型", "大小", "无法解析原因"]
        for idx, head in enumerate(headers):
            _set_cell_text(table.rows[0].cells[idx], head, size=9.5, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_TITLE)
            _shade_cell(table.rows[0].cells[idx], SHADE_HEADER)
        _set_repeat_header(table.rows[0])

        for idx, res in enumerate(bad, start=1):
            cells = table.add_row().cells
            _set_cell_text(cells[0], str(idx), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(cells[1], res.file_name, size=9)
            _set_cell_text(cells[2], res.file_type, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(cells[3], res.size_text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(cells[4], res.error_message, size=9, color=COLOR_HIGH)
        _set_col_widths(table, [1.3, 5.2, 1.8, 2.0, 4.7])

    # ------------------------------------------------------------------
    def _suggestions(self, doc: Any) -> None:
        """六、整改建议。"""
        _h1(doc, "六、整改建议")
        by_rule: Dict[str, int] = dict(self.summary.get("by_rule", {}) or {})
        sev = self.summary.get("severity", {}) or {}

        _h2(doc, "6.1 总体整改要求")
        general = [
            "按「严重 → 一般 → 轻微」的顺序推进整改，严重问题应在归档前全部清零。",
            "整改应在原始文档副本上进行，保留修改痕迹与版本记录，便于复核与追溯。",
            "整改完成后使用本工具重新检测，确认问题闭环后再进入归档或报送流程。",
            "涉密文档整改全过程应在本机离线环境完成，禁止通过网络传输或外发。",
        ]
        for idx, text in enumerate(general, start=1):
            _para(doc, f"{idx}. {text}", size=10.5, line_spacing=1.5, space_after=3)

        if by_rule:
            _para(doc, "", space_after=4)
            _h2(doc, "6.2 高频问题针对性建议")
            top_rules = list(by_rule.items())[:8]
            suggestion_map = self._collect_suggestions()
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            headers = ["序号", "问题类型", "次数", "整改建议"]
            for idx, head in enumerate(headers):
                _set_cell_text(table.rows[0].cells[idx], head, size=9.5, bold=True,
                               align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_TITLE)
                _shade_cell(table.rows[0].cells[idx], SHADE_HEADER)
            _set_repeat_header(table.rows[0])
            for idx, (name, count) in enumerate(top_rules, start=1):
                cells = table.add_row().cells
                _set_cell_text(cells[0], str(idx), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(cells[1], name, size=9)
                _set_cell_text(cells[2], str(count), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(cells[3], suggestion_map.get(name, "请按规范人工复核修正。"), size=9)
            _set_col_widths(table, [1.3, 4.2, 1.5, 8.0])

        _para(doc, "", space_after=4)
        _h2(doc, "6.3 结论")
        if int(sev.get("high", 0)) > 0:
            concl = ("本次检测发现严重级别问题，相关文档暂不具备归档 / 报送条件，"
                     "请优先完成严重问题整改并重新检测。")
        elif int(self.summary.get("total_issues", 0)) > 0:
            concl = ("本次检测未发现严重级别问题，但存在一般与轻微问题，"
                     "建议在正式归档 / 报送前完成规范化整理。")
        else:
            concl = "本次检测未发现低级错误，文档规范性良好，具备归档 / 报送条件。"
        _para(doc, concl, size=10.5, bold=True, line_spacing=1.5, first_line_indent=21)

    def _collect_suggestions(self) -> Dict[str, str]:
        """从检测结果中提取 规则名 -> 建议 的映射。"""
        out: Dict[str, str] = {}
        for res in self.results:
            for issue in res.issues:
                if issue.rule_title not in out and issue.suggestion:
                    out[issue.rule_title] = issue.suggestion
        return out

    # ------------------------------------------------------------------
    def _footer_note(self, doc: Any) -> None:
        """报告落款与保密声明。"""
        _para(doc, "", space_after=20)
        _para(doc, "— 报告结束 —", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=COLOR_MUTED, space_after=16)
        _para(doc, f"报告生成时间：{self.now.strftime('%Y-%m-%d %H:%M:%S')}",
              size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_MUTED, space_after=2)
        _para(doc, "生成工具：文档低级错误检查工具（本地离线版）",
              size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_MUTED, space_after=2)
        _para(doc, "保密声明：本报告及检测过程全程离线，未产生任何网络传输记录。",
              size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_MUTED, space_after=0)


def build_report(results: List[FileResult], summary: Dict[str, Any], out_path: str,
                 rule_summary: Optional[Dict[str, Any]] = None,
                 operator: str = "", org: str = "",
                 include_cover: bool = True, detail_columns: Optional[List[str]] = None) -> str:
    """函数式入口：生成 Word 报告并返回保存路径。"""
    builder = ReportBuilder(results, summary, rule_summary, operator, org,
                            include_cover=include_cover, detail_columns=detail_columns)
    return builder.build(out_path)


def default_report_name(prefix: str = "文档低级错误检查报告") -> str:
    """生成默认报告文件名（含时间戳，避免覆盖）。"""
    return f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
