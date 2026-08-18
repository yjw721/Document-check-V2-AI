# -*- coding: utf-8 -*-
"""
Word (.docx) 低级错误检测引擎
=========================================
覆盖检测项：
    1. 多余全角空格 / 段首冗余空格 / 段尾空格 / 无效空白字符 / 连续多空格
    2. 中英文标点混用 / 连续重复标点 / 括号引号不配对
    3. 空标题 / 连续空段落 / 大量无效换行
    4. 自动编号断裂 / 手工序号不连续 / 序号格式混乱
    5. 表格空白行 / 空白列 / 跨页无重复表头 / 表头空单元格
    6. 图片对象加载异常 / 空白冗余内容

实现方式：
    python-docx 解析 + 底层 XML(w:numPr / w:br / w:tblHeader) 读取，
    全部本地计算，绝不联网、绝不修改原始文档（只读打开）。
"""

from __future__ import annotations

import os
import re
import zipfile
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from checkers.base import (
    STATUS_ISSUE,
    STATUS_PASS,
    STATUS_UNREADABLE,
    FileResult,
    Issue,
    clip,
    precheck_ooxml,
    visualize_ws,
    word_location,
)
from config.config_manager import RuleConfig
from checkers.textnorm_checker import TextNormChecker
from checkers.fluency_checker import FluencyChecker
from checkers.custom_rules import CustomRuleEngine
from checkers.wordbank import WordBankEngine

# ---------------------------------------------------------------------------
# 正则与字符集常量
# ---------------------------------------------------------------------------
FULL_WIDTH_SPACE = "\u3000"
# 无效/异常空白字符：不换行空格、零宽系列、BOM、垂直制表、换页
INVALID_WS_CHARS = ["\u00a0", "\u200b", "\u200c", "\u200d", "\ufeff", "\u000b", "\u000c", "\u2028", "\u2029"]

CJK_RE = re.compile(r"[\u4e00-\u9fff\u3400-\u4dbf]")
# 中文语境里出现的英文标点（前后至少一侧为中文）
EN_PUNCT_IN_CJK_RE = re.compile(
    r"(?:(?<=[\u4e00-\u9fff])[,;:!?]|[,;:!?](?=[\u4e00-\u9fff])"
    r"|(?<=[\u4e00-\u9fff])\.(?![0-9a-zA-Z])|(?<=[\u4e00-\u9fff])\((?=[^)]*[\u4e00-\u9fff])"
    r"|(?<=[\u4e00-\u9fff])\))"
)
# 连续重复标点（中英文）
REPEAT_PUNCT_RE = re.compile(r"([，。；：！？、,;:!?])\1+")
# 连续多个半角空格
MULTI_SPACE_RE = re.compile(r"[ ]{2,}")

# 手工序号模式：1. / 1、/ (1) / （1）/ 一、/ 第一条
NUM_PATTERNS: List[Tuple[str, re.Pattern, str]] = [
    ("arabic_dot", re.compile(r"^\s*(\d{1,3})\s*[.．]\s*(?![0-9])"), "1. 形式"),
    ("arabic_pause", re.compile(r"^\s*(\d{1,3})\s*[、]"), "1、形式"),
    ("arabic_paren", re.compile(r"^\s*[（(]\s*(\d{1,3})\s*[）)]"), "（1）形式"),
    ("cn_pause", re.compile(r"^\s*([一二三四五六七八九十百]{1,4})\s*[、]"), "一、形式"),
    ("cn_article", re.compile(r"^\s*第\s*([一二三四五六七八九十百零〇\d]{1,5})\s*条"), "第一条 形式"),
]

CN_NUM_MAP = {
    "零": 0, "〇": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}

# 成对符号（括号 / 引号 / 书名号）
PAIRS: List[Tuple[str, str, str]] = [
    ("（", "）", "全角圆括号"),
    ("(", ")", "半角圆括号"),
    ("【", "】", "方头括号"),
    ("《", "》", "书名号"),
    ("“", "”", "中文双引号"),
    ("‘", "’", "中文单引号"),
    ("[", "]", "半角方括号"),
]


def _cn_to_int(text: str) -> Optional[int]:
    """把简单中文数字转成整数（支持 一 ~ 九十九、以及阿拉伯数字串）。"""
    text = text.strip()
    if text.isdigit():
        try:
            return int(text)
        except ValueError:
            return None
    if not text:
        return None
    if text == "十":
        return 10
    # 形如 十二 / 二十 / 二十三
    if "十" in text:
        left, _, right = text.partition("十")
        tens = CN_NUM_MAP.get(left, 1) if left else 1
        ones = CN_NUM_MAP.get(right, 0) if right else 0
        if left and left not in CN_NUM_MAP:
            return None
        if right and right not in CN_NUM_MAP:
            return None
        return tens * 10 + ones
    if len(text) == 1:
        return CN_NUM_MAP.get(text)
    return None


class WordChecker:
    """Word 文档检测器。一个实例对应一次文件检测。"""

    def __init__(self, config: RuleConfig, fluency_sensitivity: str = "normal",
                 progress: Optional[Callable[[float, str, str], bool]] = None) -> None:
        self.cfg = config
        self.kind = "word"
        self.issues: List[Issue] = []
        self._limit = config.max_issues_per_file()
        self._fluency_sensitivity = fluency_sensitivity
        self._progress = progress
        # 页码定位状态（§2）：行数估算 + 显式分页符检测
        self._page = 1
        self._lines = 0
        self._break_found = False

    # ------------------------------------------------------------------
    # 进度钩子（阶段百分比 + 思考日志；返回 False 表示任务已取消）
    # ------------------------------------------------------------------
    def _hook(self, pct: float, stage: str, log: str) -> bool:
        if not self._progress:
            return True
        try:
            return bool(self._progress(pct, stage, log))
        except Exception:  # noqa: BLE001 - 钩子异常不影响检测
            return True

    # ------------------------------------------------------------------
    # 内部工具
    # ------------------------------------------------------------------
    def _on(self, rule_key: str) -> bool:
        """规则开关判断 + 问题数量上限保护。"""
        if len(self.issues) >= self._limit:
            return False
        return self.cfg.is_enabled(self.kind, rule_key)

    def _add(self, rule_key: str, location: str, detail: str, snippet: str = "") -> None:
        """记录一条问题。"""
        if len(self.issues) >= self._limit:
            return
        self.issues.append(
            Issue(
                rule_key=rule_key,
                rule_title=self.cfg.title(self.kind, rule_key),
                severity=self.cfg.severity(self.kind, rule_key),
                location=location,
                detail=detail,
                snippet=clip(snippet, 100),
                suggestion=self.cfg.suggestion(self.kind, rule_key),
            )
        )

    # ------------------------------------------------------------------
    # 主入口
    # ------------------------------------------------------------------
    def check(self, path: str, display_name: Optional[str] = None) -> FileResult:
        """
        检测单个 .docx 文件，返回 FileResult。

        任何解析异常都会被捕获并转换为「无法解析」结果，
        保证批量检测时单文件失败不中断整体任务。
        """
        name = display_name or os.path.basename(path)
        try:
            size = os.path.getsize(path)
        except OSError:
            size = 0

        result = FileResult(file_name=name, file_path=path, file_type="Word", file_size=size)

        # 1) 结构预检：加密 / 损坏 / 非 OOXML
        if not self._hook(2.0, "parse", f"正在解析「{name}」：docx OOXML 结构预检…"):
            return result
        reason = precheck_ooxml(path)
        if reason:
            result.status = STATUS_UNREADABLE
            result.error_message = reason
            return result

        # 2) 打开文档
        if not self._hook(6.0, "parse", f"正在解析「{name}」：读取 docx 文档正文与样式…"):
            return result
        try:
            doc = Document(path)
        except Exception as exc:  # noqa: BLE001 - 兜底所有解析异常
            result.status = STATUS_UNREADABLE
            result.error_message = f"文档解析失败：{type(exc).__name__} - {clip(str(exc), 120)}"
            return result

        # 3) 执行各类检测（任一检测项异常不影响其它项）
        self.issues = []
        self._page = 1
        self._lines = 0
        self._break_found = False
        # 文字规范 / 表述问题 / 自定义规则 / 自定义词库检测引擎（共享同一 issues 列表）
        self.tn = TextNormChecker(self.cfg, self.issues, self._limit)
        self.fl = FluencyChecker(self.cfg, self.issues, self._limit, self._fluency_sensitivity)
        self.cust = CustomRuleEngine(self.issues, self._limit)
        self.wb = WordBankEngine(self.issues, self._limit)
        para_stat = {"paragraphs": 0, "headings": 0, "tables": 0, "images": 0}

        try:
            paragraphs = list(doc.paragraphs)
            para_stat["paragraphs"] = len(paragraphs)
            if not self._hook(10.0, "page",
                              f"正在读取「{name}」docx 原生分页（lastRenderedPageBreak）获取页码…"):
                return result
            self._check_paragraphs(paragraphs, para_stat)
            if not self._hook(80.0, "format_error",
                              f"正在检测「{name}」编号连续性与手工序号…"):
                return result
            self._check_numbering(paragraphs)
            self._check_tail_blank(paragraphs)
        except Exception as exc:  # noqa: BLE001
            self._add_engine_note(f"段落检测过程出现异常：{type(exc).__name__}")

        try:
            tables = list(doc.tables)
            para_stat["tables"] = len(tables)
            if not self._hook(86.0, "format_error",
                              f"正在检测「{name}」表格：表头/空白行列/跨页表头…"):
                return result
            self._check_tables(tables)
        except Exception as exc:  # noqa: BLE001
            self._add_engine_note(f"表格检测过程出现异常：{type(exc).__name__}")

        try:
            if not self._hook(93.0, "format_error",
                              f"正在检测「{name}」图片与对象引用完整性…"):
                return result
            img_count = self._check_images(path, doc)
            para_stat["images"] = img_count
        except Exception as exc:  # noqa: BLE001
            self._add_engine_note(f"图片检测过程出现异常：{type(exc).__name__}")

        self._hook(100.0, "summary",
                   f"正在收集「{name}」位置信息：页码/段落号…")
        result.issues = self.issues
        result.stats = para_stat
        result.truncated = len(self.issues) >= self._limit
        result.status = STATUS_ISSUE if self.issues else STATUS_PASS
        return result

    def _add_engine_note(self, msg: str) -> None:
        """引擎内部提示（作为一条轻微问题记录，便于用户知晓）。"""
        self.issues.append(
            Issue(
                rule_key="engine_note",
                rule_title="检测引擎提示",
                severity="low",
                location="—",
                detail=msg,
                snippet="",
                suggestion="该部分内容结构特殊，建议人工重点复核。",
            )
        )

    # ------------------------------------------------------------------
    # 页码定位（§2）
    # ------------------------------------------------------------------
    def _update_page(self, para: Any, text: str) -> int:
        """根据显式分页符或行数估算更新并返回当前页码。"""
        br_qn = qn("w:br")
        lrb_qn = qn("w:lastRenderedPageBreak")
        type_qn = qn("w:type")
        if not self._break_found:
            try:
                brs = para._p.findall(".//" + br_qn)
                if any((br.get(type_qn) == "page") for br in brs):
                    self._break_found = True
                    self._page += 1
                    return self._page
                if para._p.findall(".//" + lrb_qn):
                    self._break_found = True
                    self._page += 1
                    return self._page
            except Exception:  # noqa: BLE001
                pass
        else:
            try:
                brs = para._p.findall(".//" + br_qn)
                if any((br.get(type_qn) == "page") for br in brs):
                    self._page += 1
            except Exception:  # noqa: BLE001
                pass
        if not self._break_found:
            lines = text.count("\n") + 1
            stripped = text.replace("\n", "")
            if stripped:
                lines += max(0, (len(stripped) - 1) // 36)
            self._lines += lines
            while self._lines >= 38:
                self._lines -= 38
                self._page += 1
        return self._page

    # ------------------------------------------------------------------
    # 段落级检测
    # ------------------------------------------------------------------
    def _check_paragraphs(self, paragraphs: List[Any], stat: Dict[str, int]) -> None:
        """遍历段落，执行空白、标点、空标题、空段落、手动换行等检测。"""
        empty_run = 0                       # 连续空段落计数
        max_continuous = int(self.cfg.param(self.kind, "empty_paragraph", "max_continuous", 1) or 1)
        max_breaks = int(self.cfg.param(self.kind, "manual_line_break", "max_breaks", 3) or 3)
        min_multi_space = int(self.cfg.param(self.kind, "multi_space", "min_count", 2) or 2)
        allow_ellipsis = bool(self.cfg.param(self.kind, "punct_repeat", "allow_ellipsis", True))

        for idx, para in enumerate(paragraphs, start=1):
            if not self._hook(0.0, "", ""):
                return
            text = para.text or ""
            style_name = ""
            try:
                style_name = (para.style.name or "") if para.style is not None else ""
            except Exception:  # noqa: BLE001
                style_name = ""
            is_heading = "Heading" in style_name or "标题" in style_name
            if is_heading:
                stat["headings"] += 1

            page = self._update_page(para, text)
            loc = word_location(page, idx, estimated=not self._break_found)

            total = len(paragraphs)
            pct = 12.0 + 66.0 * (idx / max(total, 1))

            # ---- 文字规范 / 表述问题 / 自定义规则 / 自定义词库（离线）----
            if not self._hook(pct, "format_error", f"正在检测格式错误：第 {idx}/{total} 段（文本规范/标点/空格）…"):
                return
            self.tn.check_text(loc, text)
            if not self._hook(pct, "fluency", f"正在检测语句通顺度：第 {idx}/{total} 段…"):
                return
            self.fl.check_text(loc, text)
            if not self._hook(pct, "wordbank", f"正在校验行业词库与自定义规则：第 {idx}/{total} 段…"):
                return
            self.cust.check_text("word", loc, text)
            self.wb.check_text("word", loc, text)

            # ---- 空标题 ----
            if is_heading and not text.strip():
                if self._on("empty_heading"):
                    self._add("empty_heading", loc,
                              f"标题样式「{style_name}」段落内容为空", "")
                empty_run = 0
                continue

            # ---- 连续空段落 ----
            if not text.strip():
                empty_run += 1
                if empty_run > max_continuous and self._on("empty_paragraph"):
                    self._add("empty_paragraph", loc,
                              f"存在连续 {empty_run} 个空段落（阈值 {max_continuous}）", "")
                # 仅含空白字符的伪内容段落
                if text and self._on("blank_redundant_content"):
                    self._add("blank_redundant_content", loc,
                              "段落仅包含空白字符（看似空行实则含空格），属冗余内容",
                              visualize_ws(text))
                continue
            empty_run = 0

            # ---- 全角空格 ----
            if FULL_WIDTH_SPACE in text and self._on("full_width_space"):
                cnt = text.count(FULL_WIDTH_SPACE)
                # 段首全角空格属于缩进习惯，单独归入段首空格；此处统计正文中间的
                inner = text.strip(FULL_WIDTH_SPACE)
                if FULL_WIDTH_SPACE in inner:
                    self._add("full_width_space", loc,
                              f"段落内含 {cnt} 处全角空格（U+3000）",
                              visualize_ws(text))

            # ---- 段首冗余空格 ----
            if self._on("leading_space"):
                lead = re.match(r"^[ \t\u3000\u00a0]+", text)
                if lead:
                    self._add("leading_space", loc,
                              f"段首存在 {len(lead.group(0))} 个手工缩进空白字符",
                              visualize_ws(text[:40]))

            # ---- 段尾空格 ----
            if self._on("trailing_space"):
                tail = re.search(r"[ \t\u3000\u00a0]+$", text)
                if tail:
                    self._add("trailing_space", loc,
                              f"段落末尾存在 {len(tail.group(0))} 个多余空白字符",
                              visualize_ws(text[-40:]))

            # ---- 无效空白字符 ----
            if self._on("invalid_whitespace"):
                hits = [c for c in INVALID_WS_CHARS if c in text]
                if hits:
                    names = "、".join(f"U+{ord(c):04X}" for c in hits)
                    self._add("invalid_whitespace", loc,
                              f"含不可见异常字符：{names}", visualize_ws(text))

            # ---- 连续多个半角空格 ----
            if self._on("multi_space"):
                m = MULTI_SPACE_RE.search(text.strip())
                if m and len(m.group(0)) >= min_multi_space:
                    self._add("multi_space", loc,
                              f"存在 {len(m.group(0))} 个连续半角空格", clip(text, 100))

            # ---- 中英文标点混用 ----
            if self._on("punct_mix") and CJK_RE.search(text):
                marks = EN_PUNCT_IN_CJK_RE.findall(text)
                if marks:
                    uniq = "、".join(sorted({m for m in marks}))
                    self._add("punct_mix", loc,
                              f"中文语境中使用英文标点：{uniq}（共 {len(marks)} 处）",
                              clip(text, 100))

            # ---- 连续重复标点 ----
            if self._on("punct_repeat"):
                for m in REPEAT_PUNCT_RE.finditer(text):
                    seq = m.group(0)
                    # 规范省略号「……」允许放行
                    if allow_ellipsis and seq.startswith("…"):
                        continue
                    self._add("punct_repeat", loc,
                              f"出现连续重复标点「{seq}」", clip(text, 100))
                    break  # 每段只报一次，避免刷屏

            # ---- 括号引号不配对 ----
            if self._on("bracket_unpaired"):
                for left, right, cname in PAIRS:
                    if left == right:
                        continue
                    lc, rc = text.count(left), text.count(right)
                    if lc != rc:
                        self._add("bracket_unpaired", loc,
                                  f"{cname} 数量不配对（左 {lc} 个 / 右 {rc} 个）",
                                  clip(text, 100))
                        break

            # ---- 段内手动换行过多 ----
            if self._on("manual_line_break"):
                try:
                    br_count = len(para._p.findall(".//" + qn("w:br")))
                except Exception:  # noqa: BLE001
                    br_count = 0
                if br_count > max_breaks:
                    self._add("manual_line_break", loc,
                              f"同一段落内含 {br_count} 个手动换行符（阈值 {max_breaks}）",
                              clip(text, 100))

    # ------------------------------------------------------------------
    # 编号检测
    # ------------------------------------------------------------------
    def _check_numbering(self, paragraphs: List[Any]) -> None:
        """自动编号断裂 + 手工序号不连续 + 序号格式混乱。"""
        # ---- 1) 自动编号断裂 ----
        if self._on("numbering_break"):
            auto_blocks: List[Tuple[int, Optional[str]]] = []   # (段号, numId)
            for idx, para in enumerate(paragraphs, start=1):
                num_id = self._auto_num_id(para)
                if num_id is not None:
                    auto_blocks.append((idx, num_id))

            # 同一 numId 的段落若被非空普通段落打断，视为编号断裂
            by_id: Dict[str, List[int]] = {}
            for idx, num_id in auto_blocks:
                by_id.setdefault(num_id, []).append(idx)

            auto_idx_set = {i for i, _ in auto_blocks}
            for num_id, idx_list in by_id.items():
                if len(idx_list) < 2:
                    continue
                for prev, curr in zip(idx_list, idx_list[1:]):
                    if curr - prev <= 1:
                        continue
                    # 中间是否存在有实际文字的非编号段落
                    broken = []
                    for mid in range(prev + 1, curr):
                        if mid in auto_idx_set:
                            continue
                        mid_text = (paragraphs[mid - 1].text or "").strip()
                        if mid_text:
                            broken.append(mid)
                    if broken:
                        self._add(
                            "numbering_break",
                            f"第 {prev} 段 → 第 {curr} 段",
                            f"自动编号列表（numId={num_id}）被第 {broken[0]} 段普通段落打断，"
                            f"存在编号重新计数风险",
                            clip(paragraphs[broken[0] - 1].text or "", 80),
                        )
                        break  # 同一编号只报一次

        # ---- 2) 手工序号连续性 + 格式一致性 ----
        check_seq = self._on("manual_number_sequence")
        check_fmt = self._on("number_format_mixed")
        if not (check_seq or check_fmt):
            return

        # 按序号「样式类型」分组收集
        collected: Dict[str, List[Tuple[int, int, str]]] = {}   # type -> [(段号, 序号值, 原文)]
        for idx, para in enumerate(paragraphs, start=1):
            text = (para.text or "").strip()
            if not text or self._auto_num_id(para) is not None:
                continue  # 自动编号不做手工序号校验
            for ptype, pattern, _label in NUM_PATTERNS:
                m = pattern.match(text)
                if not m:
                    continue
                val = _cn_to_int(m.group(1))
                if val is None:
                    break
                collected.setdefault(ptype, []).append((idx, val, text))
                break

        # 序号不连续检测
        if check_seq:
            for ptype, items in collected.items():
                if len(items) < 3:
                    continue  # 样本太少易误报
                for (i1, v1, _t1), (i2, v2, t2) in zip(items, items[1:]):
                    if v2 == v1 + 1:
                        continue
                    if v2 == 1 and v1 >= 2:
                        continue  # 视为新一组列表重新开始
                    kind_desc = {"repeat": "重复", "skip": "跳号", "back": "倒序"}
                    flag = "repeat" if v2 == v1 else ("back" if v2 < v1 else "skip")
                    self._add(
                        "manual_number_sequence",
                        f"第 {i2} 段",
                        f"手工序号{kind_desc[flag]}：上一个序号为 {v1}（第 {i1} 段），"
                        f"当前为 {v2}，编号不连续",
                        clip(t2, 80),
                    )

        # 序号格式混乱检测：阿拉伯数字系列混用多种写法
        if check_fmt:
            arabic_types = [t for t in ("arabic_dot", "arabic_pause", "arabic_paren") if t in collected]
            if len(arabic_types) >= 2:
                label_map = {t: lab for t, _p, lab in NUM_PATTERNS}
                used = "、".join(label_map[t] for t in arabic_types)
                sample_idx = min(collected[t][0][0] for t in arabic_types)
                self._add(
                    "number_format_mixed",
                    f"第 {sample_idx} 段起",
                    f"文档中混用了 {len(arabic_types)} 种阿拉伯数字序号格式：{used}",
                    clip(collected[arabic_types[0]][0][2], 80),
                )

    @staticmethod
    def _auto_num_id(para: Any) -> Optional[str]:
        """
        读取段落的自动编号 numId（无自动编号返回 None）。

        两种来源都要考虑：
            1. 段落直接格式 w:pPr/w:numPr（手工套用编号）
            2. 段落样式中定义的编号（如「List Number」列表样式）
        """
        # ---- 1) 段落直接格式 ----
        try:
            pPr = para._p.pPr
            if pPr is not None and pPr.numPr is not None and pPr.numPr.numId is not None:
                return str(pPr.numPr.numId.val)
        except Exception:  # noqa: BLE001
            pass

        # ---- 2) 段落样式（含 basedOn 继承链，最多向上追溯 5 层）----
        try:
            style = para.style
            depth = 0
            while style is not None and depth < 5:
                el = getattr(style, "element", None)
                if el is not None:
                    num_id_nodes = el.findall(
                        ".//" + qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId")
                    )
                    if num_id_nodes:
                        val = num_id_nodes[0].get(qn("w:val"))
                        if val:
                            return f"style:{val}"
                style = getattr(style, "base_style", None)
                depth += 1
        except Exception:  # noqa: BLE001
            pass
        return None

    # ------------------------------------------------------------------
    # 表格检测
    # ------------------------------------------------------------------
    def _check_tables(self, tables: List[Any]) -> None:
        """表格空白行/列、跨页表头、表头空单元格。"""
        min_rows = int(self.cfg.param(self.kind, "table_no_header_repeat", "min_rows", 12) or 12)

        for t_idx, table in enumerate(tables, start=1):
            try:
                rows = list(table.rows)
            except Exception:  # noqa: BLE001
                continue
            if not rows:
                continue

            # 估算表格所占行数（用于页码定位）
            if not self._break_found:
                self._lines += min(len(rows), 40)
                while self._lines >= 38:
                    self._lines -= 38
                    self._page += 1

            # 构建文本矩阵（合并单元格会重复出现同一文本，属正常）
            matrix: List[List[str]] = []
            for row in rows:
                try:
                    matrix.append([(c.text or "").strip() for c in row.cells])
                except Exception:  # noqa: BLE001
                    matrix.append([])

            # ---- 文字规范 / 表述问题 / 自定义规则 / 自定义词库（表格单元格文本）----
            for r_idx, cells in enumerate(matrix, start=1):
                for c_idx, val in enumerate(cells, start=1):
                    if val:
                        cell_loc = (f"第 {self._page} 页"
                                     + ("（估算）" if not self._break_found else "")
                                     + f" · 表 {t_idx} 第 {r_idx} 行第 {c_idx} 列")
                        self.tn.check_text(cell_loc, val)
                        self.fl.check_text(cell_loc, val)
                        self.cust.check_text("word", cell_loc, val)
                        self.wb.check_text("word", cell_loc, val)

            # ---- 空白行 ----
            if self._on("table_empty_row"):
                for r_idx, cells in enumerate(matrix, start=1):
                    if cells and all(not c for c in cells):
                        self._add("table_empty_row", f"表 {t_idx} 第 {r_idx} 行",
                                  f"该行 {len(cells)} 个单元格全部为空", "")

            # ---- 空白列 ----
            if self._on("table_empty_col") and matrix:
                col_count = max((len(r) for r in matrix), default=0)
                for c_idx in range(col_count):
                    col_vals = [r[c_idx] for r in matrix if c_idx < len(r)]
                    if col_vals and all(not v for v in col_vals):
                        self._add("table_empty_col", f"表 {t_idx} 第 {c_idx + 1} 列",
                                  f"该列 {len(col_vals)} 个单元格全部为空", "")

            # ---- 表头空单元格 ----
            if self._on("table_empty_header_cell") and matrix:
                header = matrix[0]
                empties = [i + 1 for i, v in enumerate(header) if not v]
                if empties and len(empties) < len(header):
                    cols = "、".join(f"第 {i} 列" for i in empties[:6])
                    self._add("table_empty_header_cell", f"表 {t_idx} 首行",
                              f"表头存在 {len(empties)} 个空单元格（{cols}）",
                              " | ".join(header[:6]))

            # ---- 跨页表格未设置重复表头 ----
            if self._on("table_no_header_repeat") and len(rows) >= min_rows:
                if not self._has_header_repeat(rows[0]):
                    self._add(
                        "table_no_header_repeat",
                        f"表 {t_idx}",
                        f"表格共 {len(rows)} 行（≥{min_rows} 行易跨页），首行未设置「标题行重复」，"
                        f"跨页后将无表头",
                        " | ".join(matrix[0][:6]) if matrix else "",
                    )

    @staticmethod
    def _has_header_repeat(first_row: Any) -> bool:
        """判断表格首行是否设置了 w:tblHeader（跨页重复标题行）。"""
        try:
            trPr = first_row._tr.find(qn("w:trPr"))
            if trPr is None:
                return False
            node = trPr.find(qn("w:tblHeader"))
            if node is None:
                return False
            val = node.get(qn("w:val"))
            return val not in ("0", "false", "off")
        except Exception:  # noqa: BLE001
            return False

    # ------------------------------------------------------------------
    # 图片 / 对象检测
    # ------------------------------------------------------------------
    def _check_images(self, path: str, doc: Any) -> int:
        """
        检测图片与嵌入对象异常：
            - 文档 XML 中引用的 rId 在关系表中缺失（断链）
            - 图片部件为 0 字节或明显损坏
            - 使用外部链接图片（离线环境必然无法显示）
        返回图片总数。
        """
        img_total = 0
        broken: List[str] = []
        enabled = self._on("image_broken")

        # 1) 统计并校验图片部件字节
        try:
            for rel_id, rel in doc.part.rels.items():
                rtype = str(rel.reltype)
                if "image" not in rtype:
                    continue
                img_total += 1
                if rel.is_external:
                    broken.append(f"{rel_id}（外部链接图片，离线环境无法显示）")
                    continue
                try:
                    blob = rel.target_part.blob
                    if blob is None or len(blob) == 0:
                        broken.append(f"{rel_id}（图片数据为 0 字节）")
                except Exception:  # noqa: BLE001
                    broken.append(f"{rel_id}（图片部件读取失败）")
        except Exception:  # noqa: BLE001
            pass

        # 2) 校验正文 XML 引用的 r:embed / r:link 是否存在于关系表
        try:
            with zipfile.ZipFile(path) as zf:
                names = set(zf.namelist())
                if "word/document.xml" in names:
                    xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
                    used_ids = set(re.findall(r'r:(?:embed|link|id)="([^"]+)"', xml))
                    rel_ids = set()
                    try:
                        rel_ids = set(doc.part.rels.keys())
                    except Exception:  # noqa: BLE001
                        pass
                    missing = sorted(i for i in used_ids if i not in rel_ids)
                    for mid in missing[:5]:
                        broken.append(f"{mid}（正文引用的关系 ID 在关系表中缺失，图片/对象断链）")
                # 3) 空字节媒体文件兜底扫描
                for nm in names:
                    if nm.startswith("word/media/"):
                        try:
                            if zf.getinfo(nm).file_size == 0:
                                broken.append(f"{os.path.basename(nm)}（媒体文件为空）")
                        except KeyError:
                            continue
        except Exception:  # noqa: BLE001
            pass

        if enabled and broken:
            uniq = list(dict.fromkeys(broken))[:8]
            self._add(
                "image_broken",
                "文档图片/嵌入对象",
                f"检测到 {len(uniq)} 处图片或对象异常：" + "；".join(uniq),
                "",
            )
        return img_total

    # ------------------------------------------------------------------
    # 文档尾部空白冗余
    # ------------------------------------------------------------------
    def _check_tail_blank(self, paragraphs: List[Any]) -> None:
        """文档结尾存在大量空段落。"""
        if not self._on("blank_redundant_content") or not paragraphs:
            return
        tail = 0
        for para in reversed(paragraphs):
            if (para.text or "").strip():
                break
            tail += 1
        if tail >= 3:
            self._add(
                "blank_redundant_content",
                f"文档结尾（末 {tail} 段）",
                f"文档末尾存在 {tail} 个连续空段落，属排版冗余",
                "",
            )


def check_word(path: str, config: RuleConfig, display_name: Optional[str] = None,
               fluency_sensitivity: str = "normal",
               progress: Optional[Callable[[float, str, str], bool]] = None) -> FileResult:
    """对外统一函数式入口。"""
    return WordChecker(config, fluency_sensitivity, progress).check(path, display_name)
