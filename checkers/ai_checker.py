# -*- coding: utf-8 -*-
"""
AI 智能核验引擎（可选模块，默认关闭）
=================================================
两种模式（由后台设置 ai 组决定）：
    local    本地 AI（Ollama 原生 /api/chat 接口，默认 http://127.0.0.1:11434）
             零联网，模型与文档数据均不出本机
    online   联网 AI（OpenAI 兼容 /chat/completions 接口，如 DeepSeek、通义、
             Kimi、OpenAI 等），需配置 base_url 与 api_key

流程：
    1. 轻量提取文件文本（Word 段落+表格 / Excel 文本单元格 / PDF 逐页）
    2. 按 max_chars 分段（受 max_requests 上限保护）
    3. 逐段调用模型，要求返回严格 JSON 数组（问题清单）
    4. 解析结果生成 Issue（rule_key=ai_verify，source=ai）追加到文件结果；
       单段失败降级为弱提醒（不中断整体核验）

保密说明：默认关闭；local 模式全程零联网；online 模式仅将待核验文本
发送至用户自行配置的接口地址，本模块不做任何其它网络行为。
"""

from __future__ import annotations

import json
import os
import re
import threading
import time
import urllib.error
import urllib.request
from typing import Any, Callable, Dict, List, Optional, Tuple

from checkers.base import Issue, clip

# 参考资料目录（用户上传的标准/词汇/规范文件，AI 核验时自动携带）
REFS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "config", "ai_refs")
REFS_META = os.path.join(REFS_DIR, "meta.json")
_REF_LOCK = threading.Lock()

# AI 核验问题条目统一规则键 / 标题
AI_RULE_KEY = "ai_verify"
AI_RULE_TITLE = "AI 智能核验"

# 全局信号量：本地/在线模型并发上限（本地模型串行最稳，避免排队打爆）
_AI_SEM = threading.Semaphore(1)

# 默认配置（与 settings 中 ai 组一致；缺省时使用）
DEFAULTS: Dict[str, Any] = {
    "enabled": False,
    "mode": "local",
    "base_url": "http://127.0.0.1:11434",
    "api_key": "",
    "model": "qwen2.5:7b",
    "timeout": 60,
    "max_chars": 3000,
    "max_requests": 10,
    "ref_enabled": True,
    "ref_max_chars": 2000,
}

# ---------------------------------------------------------------------------
# 本地模型同步：扫描 Ollama 已安装模型；配置的模型不在本机时自动切换
# ---------------------------------------------------------------------------
_model_cache: Dict[str, Any] = {"ts": 0.0, "models": [], "base": ""}
_MODEL_CACHE_TTL = 10.0


def list_local_models(base_url: str) -> List[str]:
    """扫描本地 Ollama 已安装模型（带 10 秒缓存），失败返回空列表。"""
    global _model_cache
    base = (base_url or DEFAULTS["base_url"]).rstrip("/")
    now = time.time()
    if _model_cache["base"] == base and now - _model_cache["ts"] < _MODEL_CACHE_TTL:
        return _model_cache["models"]
    try:
        req = urllib.request.Request(base + "/api/tags")
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        models = sorted(str(m.get("name") or "") for m in data.get("models", []) if m.get("name"))
    except Exception:  # noqa: BLE001 - Ollama 离线不影响调用方
        models = []
    _model_cache = {"ts": now, "models": models, "base": base}
    return models


def resolve_local_model(cfg: Dict[str, Any]) -> Tuple[str, str]:
    """本地模式模型解析：配置的 model 若不在已安装列表，自动同步为首个可用模型。

    返回 (实际模型名, 提示说明)；Ollama 离线或列表为空时保持原配置、无提示。
    """
    model = str(cfg.get("model") or DEFAULTS["model"])
    base = str(cfg.get("base_url") or DEFAULTS["base_url"])
    models = list_local_models(base)
    if not models:
        return model, ""
    if model in models:
        return model, ""
    return models[0], f"本地模型已自动同步为「{models[0]}」（原配置「{model}」不在本机）"

_SYS_PROMPT = (
    "你是一名文档质量核查专家。以下是待核查的文档片段。\n"
    "请只针对【语义与内容层面】的问题进行核查，例如：前后矛盾、表意不清、"
    "逻辑不通、内容缺失、事实或数字表述错误、语句不通顺等；\n"
    "不要检查格式排版、标点符号、错别字、单位符号等机械性、规则类问题。\n"
    "如果提供了参考资料（行业标准 / 术语定义 / 书写规范），请严格依据参考资料"
    "核查：文档表述与参考标准不符、用词不合规范、术语含义使用错误等，均视为问题；"
    "参考资料未覆盖的内容按你的通用知识判断。\n"
    "请返回严格的 JSON 数组（不要返回任何其它内容或解释），格式如下：\n"
    '[{"quote":"问题原文片段","issue":"问题类型","detail":"问题说明",'
    '"suggestion":"修改建议","severity":"high或medium或low"}]\n'
    "若片段没有问题，返回空数组 []。\n"
)


class AiError(Exception):
    """AI 调用失败（连接 / 超时 / 响应异常）。"""


def _http_json(url: str, payload: Dict[str, Any], headers: Dict[str, str],
               timeout: float) -> Dict[str, Any]:
    """发送 JSON POST 请求并解析响应（仅标准库 urllib）。"""
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={
        "Content-Type": "application/json", **headers,
    })
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
    except urllib.error.HTTPError as exc:
        body = ""
        try:
            body = exc.read().decode("utf-8", errors="replace")[:200]
        except Exception:  # noqa: BLE001
            pass
        raise AiError(f"HTTP {exc.code}：{body or exc.reason}") from exc
    except (urllib.error.URLError, TimeoutError, OSError) as exc:
        raise AiError(f"连接失败：{exc}") from exc
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise AiError(f"响应不是合法 JSON：{raw[:120]}") from exc


def _call_ollama(cfg: Dict[str, Any], messages: List[Dict[str, str]],
                 think: Optional[bool] = None,
                 options: Optional[Dict[str, Any]] = None,
                 on_token: Optional[Callable[[str, str], None]] = None) -> str:
    """Ollama 原生 /api/chat（本地，零联网）。

    think=True/False：显式启用/禁用思考链（qwen3 等 thinking 模型生效，
    其它模型忽略该字段）。
    options：透传 Ollama options（如 num_predict 限制输出长度）。
    on_token(text, kind)：可选流式回调，kind ∈ {"content", "thinking"}，
    用于实时展示本地模型推理过程；传入时以 stream 模式逐 chunk 输出。
    """
    base = str(cfg.get("base_url") or DEFAULTS["base_url"]).rstrip("/")
    url = base + "/api/chat"
    payload: Dict[str, Any] = {
        "model": cfg.get("model") or DEFAULTS["model"],
        "messages": messages,
        "stream": bool(on_token),
    }
    if think is not None:
        payload["think"] = think
    if options:
        payload["options"] = options
    if on_token:
        return _call_ollama_stream(url, payload, on_token,
                                   float(cfg.get("timeout") or DEFAULTS["timeout"]))
    resp = _http_json(url, payload, {}, float(cfg.get("timeout") or DEFAULTS["timeout"]))
    msg = resp.get("message") or {}
    content = msg.get("content")
    if not isinstance(content, str):
        raise AiError(f"Ollama 响应缺少 message.content：{str(resp)[:120]}")
    return content


def _call_ollama_stream(url: str, payload: Dict[str, Any],
                        on_token: Callable[[str, str], None], timeout: float) -> str:
    """流式调用 Ollama /api/chat：逐行解析 NDJSON chunk 并回调增量。

    chunk 结构：{"message":{"content":"...","thinking":"..."}}。
    qwen3 等 thinking 模型启用 think=True 时 thinking 字段先行输出，
    非思考模型只有 content。返回拼接后的完整 content。
    """
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    buf: List[str] = []
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            for raw in resp:
                line = raw.decode("utf-8", errors="replace").strip()
                if not line:
                    continue
                try:
                    chunk = json.loads(line)
                except json.JSONDecodeError:
                    continue
                msg = chunk.get("message") or {}
                content = msg.get("content")
                if isinstance(content, str) and content:
                    buf.append(content)
                    on_token(content, "content")
                thinking = msg.get("thinking")
                if isinstance(thinking, str) and thinking:
                    on_token(thinking, "thinking")
                if chunk.get("done"):
                    break
    except urllib.error.HTTPError as exc:
        body = ""
        try:
            body = exc.read().decode("utf-8", errors="replace")[:200]
        except Exception:  # noqa: BLE001
            pass
        raise AiError(f"HTTP {exc.code}：{body or exc.reason}") from exc
    except (urllib.error.URLError, TimeoutError, OSError) as exc:
        raise AiError(f"连接失败：{exc}") from exc
    return "".join(buf)


def _call_openai(cfg: Dict[str, Any], messages: List[Dict[str, str]]) -> str:
    """OpenAI 兼容 /chat/completions（联网）。"""
    base = str(cfg.get("base_url") or "").rstrip("/")
    if not base:
        raise AiError("未配置联网 AI 接口地址 base_url")
    url = base + "/chat/completions" if base.endswith("/v1") else base + "/v1/chat/completions"
    headers = {}
    key = str(cfg.get("api_key") or "").strip()
    if key:
        headers["Authorization"] = "Bearer " + key
    resp = _http_json(url, {
        "model": cfg.get("model") or "deepseek-chat",
        "messages": messages,
        "stream": False,
        "temperature": 0.2,
    }, headers, float(cfg.get("timeout") or DEFAULTS["timeout"]))
    try:
        return resp["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise AiError(f"OpenAI 兼容响应结构异常：{str(resp)[:120]}") from exc


def _parse_issues(content: str) -> List[Dict[str, Any]]:
    """从模型输出中提取 JSON 数组（容忍 ```json 包裹与前后缀文本）。"""
    text = content.strip()
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1 or end <= start:
        return []
    try:
        arr = json.loads(text[start:end + 1])
    except json.JSONDecodeError:
        return []
    if not isinstance(arr, list):
        return []
    out: List[Dict[str, Any]] = []
    for it in arr:
        if not isinstance(it, dict):
            continue
        quote = str(it.get("quote") or "").strip()
        if not quote:
            continue
        sev = str(it.get("severity") or "medium")
        if sev not in ("high", "medium", "low"):
            sev = "medium"
        out.append({
            "quote": clip(quote, 120),
            "issue": str(it.get("issue") or "语义问题"),
            "detail": str(it.get("detail") or "").strip(),
            "suggestion": str(it.get("suggestion") or "").strip(),
            "severity": sev,
        })
    return out


# ---------------------------------------------------------------------------
# 文本提取（轻量版：Word / Excel / PDF）
# ---------------------------------------------------------------------------
def extract_texts(path: str, ftype: str) -> List[Tuple[str, str]]:
    """提取 [(位置, 文本), ...]，任何解析失败返回空列表。"""
    try:
        if ftype == "Word":
            return _extract_word(path)
        if ftype == "Excel":
            return _extract_excel(path)
        if ftype == "PDF":
            return _extract_pdf(path)
    except Exception:  # noqa: BLE001
        pass
    return []


def _extract_word(path: str) -> List[Tuple[str, str]]:
    from docx import Document
    doc = Document(path)
    blocks: List[Tuple[str, str]] = []
    for i, p in enumerate(doc.paragraphs, start=1):
        t = (p.text or "").strip()
        if len(t) >= 4:
            blocks.append((f"第 {i} 段", t))
    for ti, table in enumerate(doc.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                t = (cell.text or "").strip()
                if len(t) >= 4:
                    blocks.append((f"表 {ti} 第 {ri} 行第 {ci} 列", t))
    return blocks


def _extract_excel(path: str) -> List[Tuple[str, str]]:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks: List[Tuple[str, str]] = []
    try:
        for ws in wb.worksheets:
            for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
                for ci, val in enumerate(row, start=1):
                    if isinstance(val, str) and len(val.strip()) >= 4:
                        blocks.append((f"「{ws.title}」第 {ri} 行第 {ci} 列", val.strip()))
    finally:
        wb.close()
    return blocks


def _extract_pdf(path: str) -> List[Tuple[str, str]]:
    from pypdf import PdfReader
    reader = PdfReader(path)
    blocks: List[Tuple[str, str]] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            t = (page.extract_text() or "").strip()
        except Exception:  # noqa: BLE001
            t = ""
        if len(t) >= 4:
            blocks.append((f"第 {i} 页", t))
    return blocks


# ---------------------------------------------------------------------------
# 参考资料（标准/词汇/规范）：上传 → 提取文本 → 核验时注入
# ---------------------------------------------------------------------------
def _refs_meta() -> List[Dict[str, Any]]:
    """读取参考资料元信息列表。"""
    try:
        if os.path.exists(REFS_META):
            with open(REFS_META, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
    except Exception:  # noqa: BLE001
        pass
    return []


def _save_refs_meta(items: List[Dict[str, Any]]) -> None:
    os.makedirs(REFS_DIR, exist_ok=True)
    with open(REFS_META, "w", encoding="utf-8") as f:
        json.dump(items, f, ensure_ascii=False, indent=1)


def list_refs() -> List[Dict[str, Any]]:
    """参考资料列表（仅保留文件仍存在的条目）。"""
    items = _refs_meta()
    out = []
    for it in items:
        name = str(it.get("name") or "")
        if not name or not os.path.exists(os.path.join(REFS_DIR, name)):
            continue
        out.append({
            "name": name,
            "chars": int(it.get("chars") or 0),
            "enabled": bool(it.get("enabled", True)),
            "updated": str(it.get("updated") or ""),
        })
    return out


def _extract_ref_text(filename: str, raw: bytes) -> str:
    """按扩展名从上传文件提取纯文本（txt/md/docx/pdf）。"""
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md", ".csv"):
        return raw.decode("utf-8", errors="replace")
    if ext == ".docx":
        import io as _io
        from docx import Document
        doc = Document(_io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)
    if ext == ".pdf":
        from pypdf import PdfReader
        import io as _io
        reader = PdfReader(_io.BytesIO(raw))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    raise ValueError(f"不支持的参考资料格式：{ext or '无扩展名'}（支持 .txt / .md / .csv / .docx / .pdf）")


def save_ref(filename: str, raw: bytes) -> Dict[str, Any]:
    """保存参考资料（提取文本为 .txt 落盘），返回条目信息。"""
    name = os.path.basename(filename or "").strip()
    if not name:
        raise ValueError("文件名不能为空")
    text = _extract_ref_text(name, raw).strip()
    if len(text) < 10:
        raise ValueError("参考资料内容过少（提取不足 10 字符），请确认文件含有效文本")
    with _REF_LOCK:
        os.makedirs(REFS_DIR, exist_ok=True)
        stem = os.path.splitext(name)[0]
        out_name = stem + ".txt"
        out_path = os.path.join(REFS_DIR, out_name)
        # 同名覆盖
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)
        items = [it for it in _refs_meta() if it.get("name") != out_name]
        items.append({
            "name": out_name,
            "chars": len(text),
            "enabled": True,
            "updated": time.strftime("%Y-%m-%d %H:%M:%S"),
        })
        _save_refs_meta(items)
    return {"name": out_name, "chars": len(text), "enabled": True}


def delete_ref(name: str) -> bool:
    """删除参考资料。"""
    name = os.path.basename(name or "").strip()
    with _REF_LOCK:
        items = _refs_meta()
        path = os.path.join(REFS_DIR, name)
        existed = os.path.exists(path)
        if existed:
            try:
                os.remove(path)
            except OSError:
                pass
        rest = [it for it in items if it.get("name") != name]
        _save_refs_meta(rest)
    return existed


def set_ref_enabled(name: str, enabled: bool) -> bool:
    """启用 / 停用参考资料。"""
    name = os.path.basename(name or "").strip()
    with _REF_LOCK:
        items = _refs_meta()
        found = False
        for it in items:
            if it.get("name") == name:
                it["enabled"] = bool(enabled)
                found = True
        if found:
            _save_refs_meta(items)
    return found


def _load_ref_text(max_chars: int) -> str:
    """合并所有启用的参考资料文本（按总长度截断）。"""
    parts: List[str] = []
    total = 0
    for it in _refs_meta():
        if not it.get("enabled", True):
            continue
        name = str(it.get("name") or "")
        path = os.path.join(REFS_DIR, name)
        if not os.path.exists(path):
            continue
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        except OSError:
            continue
        text = text.strip()
        if not text:
            continue
        head = text[:max_chars - total] if max_chars > total else ""
        if head:
            parts.append(f"【{name}】\n{head}")
            total += len(head)
        if total >= max_chars:
            break
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# 分段与核验
# ---------------------------------------------------------------------------
def _chunk_blocks(blocks: List[Tuple[str, str]], max_chars: int,
                  max_requests: int) -> List[Tuple[str, str]]:
    """把文本块按 max_chars 合并分段，最多 max_requests 段。"""
    chunks: List[Tuple[str, str]] = []
    cur_loc, cur_text = "", []
    cur_len = 0
    for loc, text in blocks:
        if not text:
            continue
        if cur_len + len(text) > max_chars and cur_text:
            chunks.append((cur_loc, "\n".join(cur_text)))
            cur_loc, cur_text, cur_len = "", [], 0
            if len(chunks) >= max_requests:
                break
        if not cur_text:
            cur_loc = loc
        cur_text.append(text)
        cur_len += len(text)
    if cur_text and len(chunks) < max_requests:
        chunks.append((cur_loc, "\n".join(cur_text)))
    return chunks[:max_requests]


def ai_check_file(path: str, ftype: str, issues_out: List[Issue],
                  cfg: Optional[Dict[str, Any]] = None,
                  limit: int = 800, max_files_issues: int = 0,
                  cancel: Optional[Callable[[], bool]] = None,
                  on_token: Optional[Callable[[str, str], None]] = None,
                  on_chunk: Optional[Callable[[int, int], None]] = None) -> Tuple[int, str]:
    """
    对单个文件执行 AI 智能核验，把命中问题追加到 issues_out。

    返回 (追加条数, 状态说明)。任何失败都不抛异常，降级为说明文字。
    cancel 可选：返回 True 表示任务已取消，请求间隙立即中止。
    on_token(text, kind)：可选流式回调（kind ∈ content / thinking），
    local 模式逐 token 触发，用于实时展示本地模型推理过程。
    on_chunk(k, total)：每段推理开始前回调（段序号 / 总段数）。
    """
    ai = {**DEFAULTS, **(cfg or {})}
    if not ai.get("enabled"):
        return 0, "AI 核验未启用"
    if cancel and cancel():
        return 0, "任务已取消"

    mode = ai.get("mode") or "local"
    sync_note = ""
    if mode == "local":
        model, sync_note = resolve_local_model(ai)
        if sync_note:
            ai["model"] = model

    blocks = extract_texts(path, ftype)
    if not blocks:
        return 0, "未能提取文本，跳过 AI 核验"

    chunks = _chunk_blocks(blocks, int(ai.get("max_chars") or 3000),
                           int(ai.get("max_requests") or 10))
    if not chunks:
        return 0, "文本为空，跳过 AI 核验"

    mode = ai.get("mode") or "local"
    ref_text = ""
    if ai.get("ref_enabled", True):
        ref_text = _load_ref_text(int(ai.get("ref_max_chars") or 2000))
    added = 0
    fail_reasons: List[str] = []
    for k, (loc, text) in enumerate(chunks, 1):
        if cancel and cancel():
            return added, f"任务已取消（已完成 {added} 条）"
        if max_files_issues > 0 and len(issues_out) >= max_files_issues:
            break
        if limit > 0 and len(issues_out) >= limit:
            break
        if on_chunk:
            try:
                on_chunk(k, len(chunks))
            except Exception:  # noqa: BLE001 - 回调异常不影响核验
                pass
        try:
            with _AI_SEM:
                content = _call_ollama(ai, _messages(text, ref_text),
                                       on_token=on_token) if mode == "local" \
                    else _call_openai(ai, _messages(text, ref_text))
            items = _parse_issues(content)
        except AiError as exc:
            fail_reasons.append(f"{loc}：{exc}")
            continue
        for it in items:
            if limit > 0 and len(issues_out) >= limit:
                break
            issues_out.append(Issue(
                rule_key=AI_RULE_KEY,
                rule_title=AI_RULE_TITLE,
                severity=it["severity"],
                location=loc,
                detail=f"{it['issue']}：{it['detail']}",
                snippet=it["quote"],
                suggestion=it["suggestion"],
                source="ai",
            ))
            added += 1

    if fail_reasons and added == 0:
        return 0, "AI 核验执行失败：" + "；".join(fail_reasons[:2])
    if fail_reasons:
        return added, "部分分段失败：" + "；".join(fail_reasons[:2])
    return added, ("完成" + (f"（{sync_note}）" if sync_note else ""))


def _messages(text: str, ref_text: str = "") -> List[Dict[str, str]]:
    msgs = [{"role": "system", "content": _SYS_PROMPT}]
    if ref_text:
        msgs.append({
            "role": "system",
            "content": "以下是用户上传的核验参考资料（行业标准 / 术语定义 / 书写规范），"
                       "核验时请严格依据这些标准检查文档，不得违背：\n" + ref_text,
        })
    msgs.append({"role": "user", "content": f"文档片段：\n{text}"})
    return msgs


def test_connection(ai: Dict[str, Any]) -> Tuple[bool, str]:
    """测试 AI 连接：发送最小请求，返回 (是否成功, 说明)。

    连通性判定：服务可达且模型返回了非空响应即视为成功
    （不同模型对“请回复 ok”的遵循度不同，不据此判失败）。
    """
    cfg = {**DEFAULTS, **(ai or {})}
    mode = cfg.get("mode") or "local"
    sync_note = ""
    if mode == "local":
        model, sync_note = resolve_local_model(cfg)
        if sync_note:
            cfg["model"] = model
    try:
        if mode == "local":
            content = _call_ollama(cfg, [{"role": "user", "content": "你好，请回复：ok"}])
        else:
            content = _call_openai(cfg, [{"role": "user", "content": "你好，请回复：ok"}])
        ok = bool((content or "").strip())
        msg = f"连接成功（模型响应：{clip(content, 60)}）" if ok \
            else "连接成功但模型未返回内容"
        if sync_note:
            msg += f"；{sync_note}"
        return ok, msg
    except AiError as exc:
        return False, str(exc)