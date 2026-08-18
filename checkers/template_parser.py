# -*- coding: utf-8 -*-
"""
范本解析生成规则引擎（本地离线、仅以范本为基准、零联网）
============================================================
职责：
    1. 批量解析用户上传的资产评估行业范本文档（.docx / .pdf），
       提取「分页文本块」(doc, page, text)
    2. 仅以上传范本内容为唯一基准，自动提取：
       - 标准表述参考（应当 / 必须 / 按照 / 依据 等规范句式）
       - 禁用用语 / 绝对化用语 / 承诺性描述（内置扫描词，仅当范本中出现才产出）
       - 口语化文字（仅当范本中出现才产出）
       - 典型错别字变体（内置错别字对，仅当范本确认正确写法才产出）
       - 规范格式（日期 / 金额 / 序号，从范本统计主格式，检测目标文档中的
         其它格式作为「格式严重错误」正则规则）
       - 多范本术语表述冲突（高频短语相似度比对 + 优选建议）
    3. 产出可直接预览、选择性确认导入的自定义规则草案与词库词条草案。
       所有建议文案引用范本原文与页码，绝不凭空编造规范要求。

风险分级标签（统一）：
    笔误警示 / 表述优化建议 / 格式严重错误 / 执业风险警示
    （绝对化用语、承诺性描述强制标记【执业风险警示】）

保密说明：全程本地内存计算，无任何网络行为、不上传任何文档内容。
"""

from __future__ import annotations

import csv
import difflib
import io
import os
import re
import struct
import uuid
from typing import Any, Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# 统一风险分级标签
# ---------------------------------------------------------------------------
TAG_TYPO = "笔误警示"
TAG_STYLE = "表述优化建议"
TAG_FORMAT = "格式严重错误"
TAG_RISK = "执业风险警示"

# 标签 → 严重级别
TAG_SEVERITY = {
    TAG_TYPO: "medium",
    TAG_STYLE: "low",
    TAG_FORMAT: "high",
    TAG_RISK: "high",
}

# 标签 → 规则组类别
TAG_CATEGORY = {
    TAG_TYPO: "expression",
    TAG_STYLE: "expression",
    TAG_FORMAT: "format_error",
    TAG_RISK: "expression",
}

# 支持的基准文件格式（文档类 + 文本类 + 搜狗词库）
_VALID_EXT = {".docx", ".pdf", ".txt", ".csv", ".scel"}

# 外部基准文件类别（7 类）：决定导入后的分组名称与无标签信息时的默认分级倾向
CATEGORY_NAMES = {
    "general": "通用检测规则库",
    "industry": "行业规范词库",
    "asset": "资产评估准则",
    "practice": "执业规范模板",
    "correction": "标准纠错库",
    "forbidden": "禁用词库",
    "official": "公文规范库",
}
CATEGORY_DEFAULT_TAG = {
    "general": TAG_TYPO,
    "industry": TAG_TYPO,
    "asset": TAG_TYPO,
    "practice": TAG_TYPO,
    "correction": TAG_TYPO,
    "forbidden": TAG_RISK,   # 禁用词库默认执业风险警示
    "official": TAG_FORMAT,  # 公文规范库默认格式严重错误
}

# 行级文本解析：纠错对分隔符 / 正则特征 / 注释
_PAIR_SEP = re.compile(r"\s*(?:=>|->|→|＝|==|，|,|\t)\s*")
_REGEX_LINE = re.compile(r"^\s*/.*/\s*$|\\d\{|\\[a-zA-Z]\{|\^|\$|[\[\]{}]\(|\(\?:|[a-zA-Z]+\(")
_COMMENT_LINE = re.compile(r"^\s*(?:#|//|;|//|--)")


def gen_id(prefix: str = "tp") -> str:
    return prefix + uuid.uuid4().hex[:8]


# ---------------------------------------------------------------------------
# 内置扫描词库（仅作为「扫描触发器」；最终是否产出以范本内容为准）
# ---------------------------------------------------------------------------
# 绝对化用语 / 承诺性描述（资产评估报告高风险表述）
RISK_WORDS = [
    "保证", "承诺", "确保", "一定", "绝对", "最优", "最佳", "唯一",
    "必然", "稳赚", "无风险", "零风险", "百分百", "百分之百", "全部责任",
    "无条件", "第一品牌", "领先水平", "国际一流", "国内首创", "绝无仅有",
    "万无一失", "板上钉钉", "必定", "毫无疑问",
]

# 口语化 / 非书面化文字
COLLOQUIAL_WORDS = [
    "咱们", "差不多", "大概", "好像", "挺不错", "蛮好", "有点", "有点儿",
    "搞一下", "弄一下", "就是说", "等等吧", "回头", "挺好的", "一大堆",
    "说白了", "说白了就是", "打个比方", "众所周知",
]

# 常见错别字 / 术语误写对：(标准写法, 常见错误写法)
TYPO_PAIRS = [
    ("截至", "截止"), ("账号", "帐号"), ("其他", "其它"),
    ("作为", "做为"), ("部署", "布署"), ("即使", "既使"),
    ("青睐", "亲睐"), ("一如既往", "一如继往"), ("坐落", "座落"),
    ("综合", "综和"), ("评估", "评诂"), ("核算", "合算"),
    ("审计", "审记"), ("分录", "分路"), ("报表", "爆表"),
    ("预算", "予算"), ("折旧", "折旧"), ("摊销", "摊消"),
    ("资产", "资産"), ("负债", "负馈"), ("权益", "权宜"),
    ("收益", "收溢"), ("利润", "利闰"), ("现金流量", "现金流动量"),
    ("净值", "净直"), ("公允", "工允"), ("持续经营", "持继经营"),
]

# 日期 / 金额 / 序号 格式模板
DATE_FORMATS = [
    ("中文", re.compile(r"\d{4}年\d{1,2}月\d{1,2}日")),
    ("短横线", re.compile(r"\d{4}-\d{1,2}-\d{1,2}")),
    ("斜杠", re.compile(r"\d{4}/\d{1,2}/\d{1,2}")),
    ("点号", re.compile(r"\d{4}\.\d{1,2}\.\d{1,2}")),
]
# 目标文档中检测「与范本主格式不一致」的其它日期格式
DATE_FORMAT_MAP = {
    "中文": (r"\d{4}[-/\.]\d{1,2}[-/\.]\d{1,2}", "范本统一使用「YYYY年M月D日」中文日期格式"),
    "短横线": (r"\d{4}年\d{1,2}月\d{1,2}日|\d{4}/\d{1,2}/\d{1,2}|\d{4}\.\d{1,2}\.\d{1,2}", "范本统一使用「YYYY-MM-DD」日期格式"),
    "斜杠": (r"\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2}|\d{4}\.\d{1,2}\.\d{1,2}", "范本统一使用「YYYY/MM/DD」日期格式"),
    "点号": (r"\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2}|\d{4}/\d{1,2}/\d{1,2}", "范本统一使用「YYYY.MM.DD」日期格式"),
}

AMOUNT_BAD_PATTERNS = [
    (r"\d+(?:\.\d+)?\s+元", "数字与「元」之间夹有空格"),
    (r"\d+(?:\.\d+)?\s+万元", "数字与「万元」之间夹有空格"),
    (r"元人民币", "「元人民币」语序不当，规范写法为「人民币…元」"),
    (r"人民币\s*(\d+(?:\.\d+)?)\s*[RMB￥¥]?\s*(?!元)", "人民币金额缺少计量单位「元」"),
]

# 行首序号风格
SEQ_STYLES = [
    ("一、", re.compile(r"^\s*[一二三四五六七八九十百]{1,4}\s*、")),
    ("1.", re.compile(r"^\s*\d{1,3}\s*[.．](?![0-9])")),
    ("1、", re.compile(r"^\s*\d{1,3}\s*、")),
    ("（1）", re.compile(r"^\s*[（(]\s*\d{1,3}\s*[）)]")),
]

# 禁止性句式动词（范本中「不得/严禁…」即范本原生约束）
FORBIDDEN_VERBS = ["不得", "严禁", "禁止", "不应", "不能"]
# 规范句式动词（提取标准表述参考）
NORM_VERBS = ["应当", "必须", "应", "按照", "依据", "根据", "遵循"]

_SENT_SPLIT_RE = re.compile(r"[。！？；\n\r]+")


# ---------------------------------------------------------------------------
# 文本提取层（Word / PDF → 分页文本块）
# ---------------------------------------------------------------------------
def _estimate_word_pages(texts: List[str]) -> List[int]:
    """按行数估算 Word 段落页码（与检测器同一估算口径：约 38 行 / 页）。"""
    pages: List[int] = []
    page = 1
    lines = 0
    for t in texts:
        n = t.count("\n") + 1
        stripped = t.replace("\n", "")
        if stripped:
            n += max(0, (len(stripped) - 1) // 36)
        lines += n
        while lines >= 38:
            lines -= 38
            page += 1
        pages.append(page)
    return pages


def extract_docx(data) -> List[Tuple[int, str]]:
    """读取 .docx（文件路径或 BytesIO）段落 + 表格文本，返回 [(估算页码, 文本), ...]。"""
    from docx import Document

    if hasattr(data, "seek"):
        data.seek(0)
    doc = Document(data)
    texts: List[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            texts.append(t)
    for table in doc.tables:
        seen: set = set()
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t and t not in seen:
                    seen.add(t)
                    texts.append(t)
    pages = _estimate_word_pages(texts)
    return [(p, t) for p, t in zip(pages, texts) if t]


def extract_pdf(data) -> List[Tuple[int, str]]:
    """读取 .pdf（文件路径或 BytesIO）逐页文本，返回 [(页码, 文本), ...]。"""
    from pypdf import PdfReader

    if hasattr(data, "seek"):
        data.seek(0)
    reader = PdfReader(data)
    out: List[Tuple[int, str]] = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            t = (page.extract_text() or "").strip()
        except Exception:  # noqa: BLE001
            t = ""
        if t:
            out.append((idx, t))
    return out


def _decode_bytes(data) -> str:
    """按 UTF-8 / GBK / UTF-16 顺序尝试解码文本类文件内容。"""
    if hasattr(data, "seek"):
        data.seek(0)
        raw = data.read()
    else:
        raw = data
    for enc in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return raw.decode("utf-8", errors="ignore")


def _plausible_scel_word(w: str) -> bool:
    """候选词合理性校验：以中文字符为主，允许少量 ASCII 字母数字（如 B股/T+0）。"""
    if not w:
        return False
    if len(w) > 64:
        return False
    cjk = sum(1 for c in w if "\u4e00" <= c <= "\u9fff")
    asc = sum(1 for c in w if c.isascii() and (c.isalnum() or c in "·_—"))
    total = len(w)
    if total == 1:
        return asc == 1 and cjk == 0
    return (cjk + asc) >= max(2, int(total * 0.6))


def _parse_dcs_record(raw: bytes, size: int, pos: int):
    """从 pos 尝试解析一条 DCS 变体词条记录。

    两种记录类型（复用拼音的精简记录会省略拼音头）：
      - 完整记录: [flag u16][plen u16][拼音区 plen B][wlen u16][词(UTF-16LE)][0a 00][freq u16][0 填充]
      - 精简记录: [wlen u16][词(UTF-16LE)][0a 00][freq u16][0 填充]
    判别：完整记录的 flag 恒为 1~4、plen 为 4~512 的小偶数；精简记录的首 u16 是词字节长
    （偶数），次 u16 是词首字码位（中文恒 >= 0x4E00，不可能是合法 plen）。二者互斥。
    成功返回 (词, 词尾偏移)，失败返回 None。
    """
    n, plen = struct.unpack_from("<HH", raw, pos)
    # 完整记录（先试）
    if n in (1, 2, 3, 4) and 4 <= plen <= 512 and plen % 2 == 0:
        wpos = pos + 4 + plen
        if wpos + 4 <= size:
            wlen = struct.unpack_from("<H", raw, wpos)[0]
            if 2 <= wlen <= 256 and wlen % 2 == 0 and wpos + 2 + wlen <= size:
                cand = raw[wpos + 2:wpos + 2 + wlen].decode("utf-16-le", errors="ignore")
                if _plausible_scel_word(cand):
                    return cand.strip(), wpos + 2 + wlen
    # 精简记录（词首 u16 即 wlen）
    wlen = n
    if 2 <= wlen <= 256 and wlen % 2 == 0 and pos + 2 + wlen <= size:
        cand = raw[pos + 2:pos + 2 + wlen].decode("utf-16-le", errors="ignore")
        if _plausible_scel_word(cand):
            return cand.strip(), pos + 2 + wlen
    return None


def _parse_dcs_variant(raw: bytes) -> List[str]:
    """解析 DCS 变体 SCEL（魔数 "DCS"）。

    布局：头部 + 拼音索引区 + 拼音表（[len u16][字母 UTF-16LE][idx u16] 三元组）+ 连续词条流。
    词条流从拼音表末尾开始，到文件末尾为止；记录为完整/精简双类型，零填充长度可变
    （跳过连续 0x00 直至下一条记录头）。起点动态定位：首个能持续解析 >= 50 条记录的
    2 字节对齐偏移（无命中时取解析条数最多者）。
    """
    size = len(raw)
    if size < 0x100:
        return []

    def _try(start: int, limit: int):
        """从 start 解析，返回 (记录数, 是否到达 limit)。"""
        pos, cnt = start, 0
        while pos + 4 <= size and cnt < limit:
            while pos < size and raw[pos] == 0:
                pos += 1
            if pos + 4 > size:
                break
            r = _parse_dcs_record(raw, size, pos)
            if r is None:
                break
            _word, pos = r
            if raw[pos:pos + 2] != b"\x0a\x00":
                break
            pos += 2
            if pos + 2 > size:
                break
            pos += 2  # 跳过 freq（2B），后续 0 填充由循环头统一跳过
            cnt += 1
        return cnt, cnt >= limit

    # 1) 动态定位起点
    best_start, best_cnt = None, 0
    scan_end = min(size - 4, 0x20000)  # 拼音区一般远小于 128KB
    for off in range(0x100, scan_end, 2):
        cnt, full = _try(off, 50)
        if full:
            best_start, best_cnt = off, cnt
            break
        if cnt > best_cnt:
            best_start, best_cnt = off, cnt
    if best_start is None or best_cnt < 10:
        return []

    # 2) 从定位起点完整解析到 EOF
    words: List[str] = []
    pos = best_start
    while pos + 4 <= size:
        while pos < size and raw[pos] == 0:
            pos += 1
        if pos + 4 > size:
            break
        r = _parse_dcs_record(raw, size, pos)
        if r is None:
            break
        word, pos = r
        if raw[pos:pos + 2] != b"\x0a\x00":
            break
        pos += 2
        if pos + 2 > size:
            break
        pos += 2
        if word:
            words.append(word)
    return words


def parse_scel_bytes(data) -> List[str]:
    """解析搜狗 SCEL 细胞词库（二进制），返回词条列表。

    兼容两种格式：
      - 标准 SCEL（魔数 "SCEL"）：0x40 文件头 + 拼音索引区(0x1538 + 0x2628) + 词条区(0x3B60)。
        词条：2B 拼音长度 + 2B 词字节长 + 词(UTF-16LE) + 扩展区(词频/拼音等)。
      - DCS 变体（魔数 "DCS"，如第三方导出/转换工具产物）：拼音表后单条连续词条流，
        完整/精简双类型记录（见 _parse_dcs_variant）。
    """
    if hasattr(data, "seek"):
        data.seek(0)
        raw = data.read()
    else:
        raw = data
    if len(raw) < 0x40:
        return []
    # DCS 变体：魔数 "DCS"（偏移 0x04 起 44 43 53）
    if raw[4:7] == b"DCS":
        return _parse_dcs_variant(raw)
    # 标准 SCEL
    if len(raw) < 0x3B60 + 12:
        return []
    words: List[str] = []
    off = 0x3B60
    try:
        while off + 12 <= len(raw):
            py_len, word_len = struct.unpack_from("<HH", raw, off)
            off += 4
            if py_len == 0 or word_len == 0 or word_len > 400:
                break
            word = raw[off:off + word_len].decode("utf-16-le", errors="ignore")
            off += word_len
            # 扩展区：4B 词频 + 4B 保留 + 2B 拼音个数 + 拼音串 + 1B
            if off + 10 > len(raw):
                break
            _freq, _resv = struct.unpack_from("<II", raw, off)
            off += 8
            py_count = struct.unpack_from("<H", raw, off)[0]
            off += 2 + py_count * 2 + 1
            w = word.strip()
            if w and len(w) >= 2 and len(w) <= 64:
                words.append(w)
    except Exception:  # noqa: BLE001 - 解析到异常位置即终止
        pass
    return words


# CSV 表头列名自适应（按关键词匹配列角色）
_COL_ERR = re.compile(r"错误|错词|错别字|误写|不规范|禁用|敏感|风险词|禁忌|口语|不当|问题词", re.I)
_COL_OK = re.compile(r"^正确|^规范|^标准|建议词|替换|纠正|纠错|修正词|^正词", re.I)
_COL_REGEX = re.compile(r"正则|pattern|表达式|匹配模式|校验规则|检测规则|^规则", re.I)
_COL_TAG = re.compile(r"标签|级别|等级|severity|风险等级|分类|类别", re.I)
_COL_SUG = re.compile(r"^建议|说明|备注|整改|提示|解释|理由", re.I)
_COL_NAME = re.compile(r"^名称|规则名称|^name", re.I)


def extract_chunks(data, ftype: str) -> List[Tuple[int, str]]:
    """统一入口：按文件类型分发提取（data 为文件路径或 BytesIO）。"""
    if ftype == "Word":
        return extract_docx(data)
    if ftype == "PDF":
        return extract_pdf(data)
    return []


# ---------------------------------------------------------------------------
# 分析层
# ---------------------------------------------------------------------------
def _sentences(text: str) -> List[str]:
    """切分句子，过滤过短 / 过长 / 纯数字 / 无意义片段。"""
    out: List[str] = []
    for raw in _SENT_SPLIT_RE.split(text):
        s = raw.strip()
        s = re.sub(r"[\u3000\u00a0]+", "", s)
        if len(s) < 4 or len(s) > 80:
            continue
        if re.fullmatch(r"[\d\s.,%¥￥元万元亿%，、（）()\-]+", s):
            continue
        out.append(s)
    return out


def _clean_phrase(s: str) -> str:
    """短语清洗：去标点 / 空白，只留中文与数字字母。"""
    s = re.sub(r"[\s\u3000\u00a0]", "", s)
    s = re.sub(r"[，。；：、！？!?,.、:;()（）「」『』【】《》“”\"'\-—–·…]", "", s)
    return s.strip()


def _ngrams(texts: List[str], lo: int = 3, hi: int = 6) -> Dict[str, int]:
    """统计中文 n-gram 词频（跨文本块聚合）。"""
    freq: Dict[str, int] = {}
    joined = "".join(texts)
    if not joined:
        return freq
    for size in range(lo, hi + 1):
        for i in range(0, len(joined) - size + 1):
            gram = joined[i:i + size]
            # 仅统计纯中文片段（含数字字母的术语如「2022年」也保留）
            if re.search(r"[\u4e00-\u9fff]", gram):
                freq[gram] = freq.get(gram, 0) + 1
    return freq


class TemplateParser:
    """范本解析器：一个实例对应一次批量解析任务。"""

    def __init__(self) -> None:
        self.docs: List[Dict[str, Any]] = []
        self.rules: List[Dict[str, Any]] = []
        self.entries: List[Dict[str, Any]] = []
        self.conflicts: List[Dict[str, Any]] = []
        self.references: List[Dict[str, Any]] = []
        self._rule_keys: set = set()
        self._entry_keys: set = set()
        self._all_chunks: List[Tuple[str, int, str]] = []   # (doc_name, page, text)
        self.category: str = "general"                      # 外部基准类别
        self._cur_doc: str = ""                             # 当前正在解析的文件名

    # ---------------- 工具 ----------------
    def _add_rule(self, name: str, mode: str, pattern: str, tag: str,
                  suggestion: str, doc: str, page: int) -> None:
        key = (mode, pattern)
        if key in self._rule_keys:
            return
        if len(self.rules) >= 200:
            return
        self._rule_keys.add(key)
        self.rules.append({
            "id": gen_id("r"), "name": name, "match_mode": mode, "pattern": pattern,
            "severity": TAG_SEVERITY[tag], "tag": tag, "category": TAG_CATEGORY[tag],
            "suggestion": suggestion, "scope": "all", "source_doc": doc,
            "source_page": page, "selected": True,
        })

    def _add_entry(self, keyword: str, tag: str, suggestion: str,
                   doc: str, page: int) -> None:
        kw = keyword.strip()
        if len(kw) < 2:
            return
        if kw in self._entry_keys:
            return
        if len(self.entries) >= 300:
            return
        self._entry_keys.add(kw)
        self.entries.append({
            "id": gen_id("e"), "keyword": kw, "tag": tag,
            "suggestion": suggestion, "source_doc": doc,
            "source_page": page, "selected": True,
        })

    @staticmethod
    def _src(doc: str, page: int) -> str:
        if page:
            return f"依据基准《{doc}》第 {page} 页"
        return f"依据外部基准《{doc}》"

    # ---------------- 提取器 ----------------
    def _scan_risk_words(self, doc: str, page: int, text: str) -> None:
        """绝对化用语 / 承诺性描述 →【执业风险警示】词条。"""
        for w in RISK_WORDS:
            if w in text:
                self._add_entry(
                    w, TAG_RISK,
                    f"{self._src(doc, page)}：范本原文含「{w}」。该表述属绝对化用语 / 承诺性描述，"
                    f"在执业报告中应谨慎使用，建议人工核查是否违反执业规范。",
                    doc, page,
                )

    def _scan_colloquial(self, doc: str, page: int, text: str) -> None:
        """口语化文字 →【表述优化建议】词条。"""
        for w in COLLOQUIAL_WORDS:
            if w in text:
                self._add_entry(
                    w, TAG_STYLE,
                    f"{self._src(doc, page)}：范本原文含口语化文字「{w}」，正式执业报告建议使用书面语表述。",
                    doc, page,
                )

    def _scan_typos(self, doc: str, page: int, text: str) -> None:
        """错别字变体 →【笔误警示】词条（仅当范本出现标准写法时产出错误写法）。"""
        for good, bad in TYPO_PAIRS:
            if good in text and bad not in self._entry_keys:
                self._add_entry(
                    bad, TAG_TYPO,
                    f"{self._src(doc, page)}：范本标准写法为「{good}」。若当前文档使用「{bad}」，"
                    f"疑似错别字 / 术语误写，请核对修正。",
                    doc, page,
                )

    def _scan_forbidden(self, doc: str, page: int, text: str) -> None:
        """禁止性句式 → 提取关键短语 →【执业风险警示】keyword 规则。"""
        for sent in _sentences(text):
            m = re.search(r"(?:不得|严禁|禁止|不应|不能)\s*([^，。；,!！?？、\n]{2,16})", sent)
            if not m:
                continue
            phrase = _clean_phrase(m.group(1))
            if len(phrase) < 2 or phrase in ("的", "其", "该"):
                continue
            self._add_rule(
                f"范本禁用表述「{phrase}」", "keyword", phrase, TAG_RISK,
                f"{self._src(doc, page)}：范本原文「{sent[:50]}…」明确禁止该表述，"
                f"命中时请人工核查是否存在违规承诺 / 违规表述。",
                doc, page,
            )

    # ---------------- 文本类基准解析（TXT / CSV / SCEL / 正则 / 纠错库） ----------------
    def _heuristic_tag(self, keyword: str) -> str:
        """内容启发式标签：绝对化 / 承诺 → 执业风险警示；口语化 → 表述优化建议；其余按类别默认。"""
        for w in RISK_WORDS:
            if w in keyword:
                return TAG_RISK
        for w in COLLOQUIAL_WORDS:
            if w in keyword:
                return TAG_STYLE
        return CATEGORY_DEFAULT_TAG.get(self.category, TAG_TYPO)

    @staticmethod
    def _tag_from_text(s: str):
        """把 CSV 标签列的中文描述映射为统一标签。"""
        if any(k in s for k in ("风险", "禁用", "绝对", "承诺", "警示", "敏感")):
            return TAG_RISK
        if any(k in s for k in ("格式", "严重", "错误")):
            return TAG_FORMAT
        if any(k in s for k in ("优化", "表述", "口语", "书面")):
            return TAG_STYLE
        if any(k in s for k in ("笔误", "错别字", "纠错", "错词", "误写")):
            return TAG_TYPO
        return None

    def _add_text_entry(self, keyword: str, suggestion: str = "", tag: str = "", cap: int = 300) -> None:
        """文本类词条（无页码定位，来源为外部基准文件）。

        cap：草案上限。规则/纠错类文本来源默认 300（需人工逐条确认）；
        SCEL 词库类来源按词库规模放大（_parse_scel 传入 10000），避免大词库被截断。
        """
        kw = keyword.strip().strip('"\'“”‘’')
        if len(kw) < 2 or len(kw) > 64:
            return
        if kw in self._entry_keys or len(self.entries) >= cap:
            return
        self._entry_keys.add(kw)
        t = tag or self._heuristic_tag(kw)
        self.entries.append({
            "id": gen_id("e"), "keyword": kw, "tag": t,
            "suggestion": suggestion or f"{self._src(self._cur_doc, 0)}：该词条应遵循统一规范表述。",
            "source_doc": self._cur_doc, "source_page": 0, "selected": True,
        })

    def _add_text_rule(self, name: str, pattern: str, suggestion: str = "") -> None:
        """文本类正则规则（外部基准文件）。"""
        tag = CATEGORY_DEFAULT_TAG.get(self.category, TAG_FORMAT)
        if tag not in (TAG_FORMAT, TAG_RISK):
            tag = TAG_FORMAT
        key = ("regex", pattern)
        if key in self._rule_keys or len(self.rules) >= 200:
            return
        self._rule_keys.add(key)
        self.rules.append({
            "id": gen_id("r"), "name": name, "match_mode": "regex", "pattern": pattern,
            "severity": TAG_SEVERITY[tag], "tag": tag, "category": TAG_CATEGORY[tag],
            "suggestion": suggestion or f"{self._src(self._cur_doc, 0)}：命中该正则的表述与外部基准规范不一致，请核查修正。",
            "scope": "all", "source_doc": self._cur_doc, "source_page": 0, "selected": True,
        })

    def _parse_text_lines(self, doc: str, text: str) -> None:
        """逐行解析 TXT / 行业纠错文本库 / 自定义正则规则文本。"""
        for raw in text.splitlines():
            line = raw.strip()
            if not line or _COMMENT_LINE.match(line) or len(line) > 200:
                continue
            # 1) 纠错对：错误=>正确 / 错误→正确 / 错误==正确
            if re.search(r"=>|->|→|＝|==", line):
                parts = [p.strip().strip('"\'“”') for p in re.split(r"=>|->|→|＝|==", line, maxsplit=1)]
                if len(parts) == 2 and all(parts) and 1 < len(parts[0]) <= 32:
                    self._add_text_entry(parts[0], f"外部纠错基准：标准写法为「{parts[1]}」", TAG_TYPO)
                    continue
            # 2) 正则规则文本：/pattern/ 或 含正则特征，支持 pattern::建议
            if _REGEX_LINE.search(line):
                pattern, _, sug = line.partition("::")
                pattern = pattern.strip()
                if pattern.startswith("/") and pattern.endswith("/") and len(pattern) > 2:
                    pattern = pattern[1:-1]
                if len(pattern) >= 3:
                    self._add_text_rule(f"外部基准正则规则", pattern, sug.strip())
                    continue
            # 3) 逗号 / 制表符分隔纠错对（错误,正确[,建议]）
            parts = [p.strip().strip('"\'“”') for p in re.split(r"[,，\t]", line, maxsplit=2)]
            if (len(parts) >= 2 and all(parts[:2])
                    and 1 < len(parts[0]) <= 32 and 1 < len(parts[1]) <= 32
                    and not _REGEX_LINE.search(parts[0])):
                bad, good = parts[0], parts[1]
                sug = parts[2] if len(parts) > 2 and parts[2] else f"外部纠错基准：标准写法为「{good}」"
                self._add_text_entry(bad, sug, TAG_TYPO)
                continue
            # 4) 普通词条
            self._add_text_entry(line)

    def _parse_csv_text(self, doc: str, text: str) -> None:
        """CSV 自适应列映射解析：正则列 → 规则；错误/正确列 → 纠错词条；单词列 → 词条。"""
        rows = [r for r in csv.reader(io.StringIO(text)) if r and any(c.strip() for c in r)]
        if not rows:
            return
        header = [h.strip() for h in rows[0]]

        def find_col(pat, default: int) -> int:
            for i, h in enumerate(header):
                if pat.search(h or ""):
                    return i
            return default

        col_err = find_col(_COL_ERR, 0)
        col_ok = find_col(_COL_OK, 1 if col_err == 0 else 0)
        col_regex = find_col(_COL_REGEX, -1)
        col_tag = find_col(_COL_TAG, -1)
        col_sug = find_col(_COL_SUG, -1)
        col_name = find_col(_COL_NAME, -1)
        has_header = any(
            _COL_ERR.search(h or "") or _COL_REGEX.search(h or "")
            or _COL_NAME.search(h or "") or _COL_TAG.search(h or "")
            for h in header
        )
        start = 1 if has_header else 0
        for row in rows[start:]:
            cells = [c.strip() for c in row]
            if not any(cells):
                continue
            if col_regex >= 0 and col_regex < len(cells) and cells[col_regex]:
                pattern = cells[col_regex]
                if len(pattern) >= 3:
                    name = (cells[col_name] if col_name >= 0 and col_name < len(cells) and cells[col_name]
                            else "外部基准正则规则")
                    sug = cells[col_sug] if col_sug >= 0 and col_sug < len(cells) else ""
                    self._add_text_rule(name, pattern, sug)
                    continue
            if col_err >= 0 and col_err < len(cells) and cells[col_err]:
                bad = cells[col_err]
                good = cells[col_ok] if col_ok >= 0 and col_ok < len(cells) else ""
                sug = cells[col_sug] if col_sug >= 0 and col_sug < len(cells) else ""
                tag = ""
                if col_tag >= 0 and col_tag < len(cells) and cells[col_tag]:
                    tag = self._tag_from_text(cells[col_tag]) or ""
                self._add_text_entry(bad, sug or (f"标准写法为「{good}」" if good else ""), tag)
                continue
            # 仅单列 → 词条
            self._add_text_entry(cells[0])

    def _parse_scel(self, doc: str, data) -> None:
        """SCEL 搜狗词库 → 词条（词库为纯词条集合，草案上限放宽至 10000，避免大词库被截断）。"""
        self._cur_doc = doc
        for w in parse_scel_bytes(data):
            self._add_text_entry(w, cap=10000, tag=CATEGORY_DEFAULT_TAG.get(self.category, TAG_TYPO))

    def _parse_text_file(self, name: str, ext: str, data) -> None:
        """文本类文件统一入口：按扩展名分发并记录解析状态。"""
        self._cur_doc = name
        if ext == ".scel":
            words = parse_scel_bytes(data)
            self.docs.append({
                "name": name, "file_type": "SCEL", "pages": 0,
                "chunks": len(words), "ok": bool(words),
                "error": "" if words else "未能解析词条（文件可能损坏或非标准 SCEL）",
            })
            for w in words:
                self._add_text_entry(w, cap=10000)
            return
        text = _decode_bytes(data)
        lines = len([l for l in text.splitlines() if l.strip()])
        if ext == ".csv":
            self._parse_csv_text(name, text)
            label = "CSV"
        else:
            self._parse_text_lines(name, text)
            label = "TXT"
        self.docs.append({
            "name": name, "file_type": label, "pages": 0,
            "chunks": lines, "ok": True, "error": "",
        })

    def _scan_date_format(self, texts: List[Tuple[str, int, str]]) -> None:
        """日期格式：范本主格式 → 检测其它格式。"""
        counts: Dict[str, int] = {}
        for _doc, _page, text in texts:
            for name, pat in DATE_FORMATS:
                counts[name] = counts.get(name, 0) + len(pat.findall(text))
        total = sum(counts.values())
        if total < 2:
            return
        main = max(counts, key=counts.get)
        # 仅当范本存在其它格式时生成规则
        others = {k: v for k, v in counts.items() if k != main and v > 0}
        if not others:
            return
        other_names = "、".join(f"「{k}」" for k in others)
        doc0, page0 = texts[0][0], texts[0][1]
        bad_re, hint = DATE_FORMAT_MAP[main]
        self._add_rule(
            f"日期格式应与范本一致（{main}格式）", "regex", bad_re, TAG_FORMAT,
            f"{self._src(doc0, page0)}：范本统一使用「{main}」日期格式，检测到其它格式"
            f"（{other_names}）。{hint}，请统一修正。",
            doc0, page0,
        )

    def _scan_amount_format(self, texts: List[Tuple[str, int, str]]) -> None:
        """金额格式：检测异常写法（仅当范本存在金额时产出）。"""
        any_amount = False
        for _doc, _page, text in texts:
            if re.search(r"\d+(?:\.\d+)?\s*(?:元|万元|亿元)", text):
                any_amount = True
                break
        if not any_amount:
            return
        for pat, desc in AMOUNT_BAD_PATTERNS:
            doc0, page0 = texts[0][0], texts[0][1]
            self._add_rule(
                f"金额写法不规范（{desc}）", "regex", pat, TAG_FORMAT,
                f"{self._src(doc0, page0)}：范本金额统一规范书写。检测到{desc}，请按范本格式修正。",
                doc0, page0,
            )

    def _scan_seq_style(self, texts: List[Tuple[str, int, str]]) -> None:
        """序号风格：范本主风格 → 检测混用其它风格。"""
        counts: Dict[str, int] = {}
        for _doc, _page, text in texts:
            for name, pat in SEQ_STYLES:
                for line in text.split("\n"):
                    if pat.match(line.strip()):
                        counts[name] = counts.get(name, 0) + 1
        total = sum(counts.values())
        if total < 3:
            return
        main = max(counts, key=counts.get)
        others = {k: v for k, v in counts.items() if k != main and v > 0}
        if not others:
            return
        doc0, page0 = texts[0][0], texts[0][1]
        other_names = "、".join(f"「{k}」" for k in others)
        self._add_rule(
            f"序号风格应与范本一致（{main}）", "regex",
            "|".join(f"(?m)^\\s*\\d{{1,3}}\\s*[、.]" if o in ("1.", "1、") else
                     f"(?m)^\\s*[（(]\\s*\\d{{1,3}}\\s*[）)]" if o == "（1）" else
                     f"(?m)^\\s*[一二三四五六七八九十百]{{1,4}}\\s*、"
                     for o in others),
            TAG_FORMAT,
            f"{self._src(doc0, page0)}：范本序号统一使用「{main}」风格，检测到混用其它风格"
            f"（{other_names}），请统一。",
            doc0, page0,
        )

    def _scan_references(self, doc: str, page: int, text: str) -> None:
        """标准表述参考（规范句式）——仅供预览参考，不生成命中规则。"""
        count = 0
        for sent in _sentences(text):
            if any(v in sent for v in NORM_VERBS) and 8 <= len(sent) <= 60:
                self.references.append({
                    "sentence": sent, "source_doc": doc, "source_page": page,
                })
                count += 1
                if count >= 20:
                    break

    def _scan_term_conflicts(self) -> None:
        """多范本术语表述冲突：高频短语相似度比对 → 冲突标记 + 优选建议。"""
        # 每文档高频短语（含中文，长度 3-6，出现 ≥2 次）
        doc_terms: List[Tuple[str, List[str]]] = []
        for doc_name in {d for d, _p, _t in self._all_chunks}:
            texts = [t for d, _p, t in self._all_chunks if d == doc_name]
            freq = _ngrams(texts, 3, 6)
            terms = [g for g, c in freq.items() if c >= 2]
            terms.sort(key=lambda g: (-freq[g], g))
            doc_terms.append((doc_name, terms[:80]))

        if len(doc_terms) < 2:
            return

        seen: set = set()
        for i in range(len(doc_terms)):
            for j in range(i + 1, len(doc_terms)):
                n1, terms1 = doc_terms[i]
                n2, terms2 = doc_terms[j]
                for a in terms1:
                    for b in terms2:
                        if a == b or len(a) != len(b):
                            continue
                        # 跳过滑动窗口平移一位的假阳性（如「评估基准日」的窗口变体「估基准日为」），
                        # 真实术语变体（如「其他应收款/其它应收款」）不受影响
                        if a[:-1] == b[1:] or a[1:] == b[:-1]:
                            continue
                        ratio = difflib.SequenceMatcher(None, a, b).ratio()
                        if ratio >= 0.8:
                            key = tuple(sorted((a, b)))
                            if key in seen:
                                continue
                            seen.add(key)
                            # 优选建议：以更常见的写法为准（无法判定频次时保留两写法）
                            suggestion = (f"范本《{n1}》与《{n2}》对同一概念写法不一致："
                                          f"「{a}」/「{b}」。建议核对后统一为一种标准写法，"
                                          f"避免同一术语混用。")
                            self.conflicts.append({
                                "topic": f"术语写法冲突：「{a}」vs「{b}」",
                                "docs": [n1, n2],
                                "statements": [
                                    {"text": a, "doc": n1},
                                    {"text": b, "doc": n2},
                                ],
                                "suggestion": suggestion,
                            })
                            # 非优选写法（b）→ 术语混用词条（表述优化建议）
                            self._add_entry(
                                b, TAG_STYLE,
                                f"{suggestion} 当前写法与另一范本「{a}」不一致，请人工核对统一。",
                                n2, 0,
                            )
                            if len(self.conflicts) >= 10:
                                return

    # ---------------- 主流程 ----------------
    def parse_files(self, files: List[Tuple[Any, str]], category: str = "general") -> Dict[str, Any]:
        """
        解析一批外部基准文件（.docx / .pdf / .txt / .csv / .scel）。

        参数: files = [(data, name), ...]，data 为 BytesIO / bytes / 文件路径
              category = 外部基准类别 key（见 CATEGORY_NAMES）
        返回: 草案 dict（docs / rules / entries / conflicts / references）
        """
        self.docs = []
        self.rules = []
        self.entries = []
        self.conflicts = []
        self.references = []
        self._rule_keys = set()
        self._entry_keys = set()
        self._all_chunks = []
        self.category = category if category in CATEGORY_NAMES else "general"

        ftype_map = {".docx": "Word", ".pdf": "PDF", ".txt": "TXT", ".csv": "CSV", ".scel": "SCEL"}
        for data, name in files:
            ext = os.path.splitext(name)[1].lower()
            ftype = ftype_map.get(ext, "")
            if not ftype:
                self.docs.append({
                    "name": name, "file_type": (ext.lstrip(".") or "?").upper(),
                    "pages": 0, "chunks": 0, "ok": False,
                    "error": f"不支持的文件格式「{ext}」，仅支持 .docx/.pdf/.txt/.csv/.scel",
                })
                continue
            try:
                if ext in (".docx", ".pdf"):
                    chunks = extract_chunks(data, ftype)
                    ok = bool(chunks)
                    self.docs.append({
                        "name": name, "file_type": ftype,
                        "pages": max((p for p, _ in chunks), default=0),
                        "chunks": len(chunks), "ok": ok,
                        "error": "" if ok else "无法提取文本（可能为扫描件 / 加密 / 损坏）",
                    })
                    if not ok:
                        continue
                    for page, text in chunks:
                        self._all_chunks.append((name, page, text))
                        self._scan_risk_words(name, page, text)
                        self._scan_colloquial(name, page, text)
                        self._scan_typos(name, page, text)
                        self._scan_forbidden(name, page, text)
                        self._scan_references(name, page, text)
                else:
                    self._parse_text_file(name, ext, data)
            except Exception as exc:  # noqa: BLE001 - 单文档失败不中断整体
                self.docs.append({
                    "name": name, "file_type": ftype, "pages": 0, "chunks": 0,
                    "ok": False, "error": f"解析失败：{type(exc).__name__} - {str(exc)[:100]}",
                })

        # 跨文档统计类提取（仅文档类范本）
        if self._all_chunks:
            self._scan_date_format(self._all_chunks)
            self._scan_amount_format(self._all_chunks)
            self._scan_seq_style(self._all_chunks)
            self._scan_term_conflicts()

        # 参考句去重截断
        seen_ref: set = set()
        refs: List[Dict[str, Any]] = []
        for r in self.references:
            k = r["sentence"]
            if k in seen_ref:
                continue
            seen_ref.add(k)
            refs.append(r)
            if len(refs) >= 60:
                break
        self.references = refs

        return self.draft()

    def draft(self) -> Dict[str, Any]:
        """返回当前草案（供前端预览 / 选择性导入）。"""
        return {
            "docs": self.docs,
            "rules": self.rules,
            "entries": self.entries,
            "conflicts": self.conflicts,
            "references": self.references,
        }

    def set_selected(self, rule_ids: List[str], entry_ids: List[str]) -> None:
        """按 id 更新选中状态（未列出的项保持原状）。"""
        rs, es = set(rule_ids or []), set(entry_ids or [])
        for r in self.rules:
            if r["id"] in rs:
                r["selected"] = True
            elif rs and r["id"] not in rs:
                # rs 非空时，未列出即视为取消选择
                r["selected"] = False
        for e in self.entries:
            if e["id"] in es:
                e["selected"] = True
            elif es and e["id"] not in es:
                e["selected"] = False

    def build_import(self, rule_ids: List[str], entry_ids: List[str]) -> Dict[str, Any]:
        """构建导入数据结构（规则组 + 词库组），返回可直接追加的配置。"""
        rs = set(rule_ids or [])
        es = set(entry_ids or [])
        rules = [r for r in self.rules if r["id"] in rs]
        entries = [e for e in self.entries if e["id"] in es]

        group_name = f"外部导入 · {CATEGORY_NAMES.get(self.category, '通用')}"
        from checkers.custom_rules import gen_id as gid_r
        from checkers.wordbank import gen_id as gid_e

        # 按标签分组建规则组（执业风险 / 格式 / 表述）
        rule_group = {
            "id": gid_r("g"), "name": f"{group_name} · {len(rules)} 条",
            "category": "expression", "scope": "all", "enabled": True,
            "rules": [{
                "id": gid_r("r"), "name": r["name"], "enabled": True,
                "match_mode": r["match_mode"], "pattern": r["pattern"],
                "severity": r["severity"], "tag": r["tag"], "suggestion": r["suggestion"],
            } for r in rules],
        }
        entry_group = {
            "id": gid_e("wb"), "name": f"{group_name} · {len(entries)} 条",
            "module": "text_word", "scope": "all", "enabled": True,
            "entries": [{
                "id": gid_e("e"), "keyword": e["keyword"], "tag": e["tag"],
                "suggestion": e["suggestion"], "enabled": True,
            } for e in entries],
        }
        return {"rule_group": rule_group, "entry_group": entry_group,
                "rule_count": len(rules), "entry_count": len(entries)}
