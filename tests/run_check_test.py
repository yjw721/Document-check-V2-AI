# -*- coding: utf-8 -*-
"""
端到端自检脚本（本地运行）
=========================================================================
执行内容：
    1. 加载本地规则配置
    2. 对 tests/samples/ 下全部样例执行检测
    3. 打印每个文件的命中情况与规则覆盖统计
    4. 生成一份 Word 检测报告，验证报告可正常写出

用法：python tests/run_check_test.py
"""

from __future__ import annotations

import os
import sys

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)

from checkers.base import STATUS_UNREADABLE  # noqa: E402
from checkers.scanner import check_single, collect_folder, summarize  # noqa: E402
from config.config_manager import load_rules, severity_label  # noqa: E402
from report.report_builder import build_report  # noqa: E402

SAMPLE_DIR = os.path.join(_ROOT, "tests", "samples")
OUT_DIR = os.path.join(_ROOT, "tests", "out")

# 仅对已知样例做自检，避免目录中遗留的临时文件（.tmp / .png 等）污染结果
SAMPLE_FILES = [
    "sample_word_bad.docx", "sample_word_good.docx", "sample_word_image_bad.docx",
    "sample_excel_bad.xlsx", "sample_excel_good.xlsx",
    "broken.docx", "legacy.doc",
    "sample_pdf_bad.pdf", "sample_pdf_scan.pdf", "sample_pdf_encrypted.pdf",
]


def main() -> int:
    config = load_rules()
    print("=" * 78)
    print("规则配置：", config.meta.get("config_name"), "版本", config.meta.get("config_version"))
    print(f"Word 规则 {config.enabled_count('word')}/{config.total_count('word')} 项，"
          f"Excel 规则 {config.enabled_count('excel')}/{config.total_count('excel')} 项，"
          f"文字规范规则 {config.enabled_count('textnorm')}/{config.total_count('textnorm')} 项")
    print("=" * 78)

    # 收集全部样例（仅已知清单，避免临时文件污染）
    files = [os.path.join(SAMPLE_DIR, n) for n in SAMPLE_FILES
             if os.path.exists(os.path.join(SAMPLE_DIR, n))]
    if not files:
        print("未找到样例文件，请先运行 tests/make_samples.py")
        return 1

    results = []
    for path in files:
        res = check_single(path, config, os.path.basename(path))
        results.append(res)

        print(f"\n▶ {res.file_name}  [{res.file_type}]  {res.size_text}  → {res.status_text}")
        if res.status == STATUS_UNREADABLE:
            print(f"   ⛔ {res.error_message}")
            continue
        if res.stats:
            print(f"   规模：{res.stats}")
        if not res.issues:
            print("   ✅ 未发现问题")
            continue

        by_rule = {}
        for issue in res.issues:
            by_rule.setdefault(issue.rule_title, []).append(issue)
        print(f"   共 {len(res.issues)} 条问题，覆盖 {len(by_rule)} 类规则：")
        for title, group in by_rule.items():
            sev = severity_label(group[0].severity)
            print(f"     · [{sev}] {title} × {len(group)}   例：{group[0].location} | "
                  f"{group[0].detail[:62]}")

    # ---- 汇总 ----
    summary = summarize(results)
    print("\n" + "=" * 78)
    print("汇总统计")
    print("=" * 78)
    print(f"文件总数 {summary['total_files']}　通过 {summary['pass_files']}　"
          f"存在问题 {summary['issue_files']}　无法解析 {summary['error_files']}")
    print(f"问题总数 {summary['total_issues']}　"
          f"严重 {summary['severity']['high']}　一般 {summary['severity']['medium']}　"
          f"轻微 {summary['severity']['low']}")
    print("\n问题类型分布：")
    for name, count in summary["by_rule"].items():
        print(f"  {count:>4} × {name}")

    # ---- 规则覆盖检查 ----
    hit_keys = {i.rule_key for r in results for i in r.issues}
    print("\n" + "=" * 78)
    print("规则命中覆盖检查")
    print("=" * 78)
    for kind in ("word", "excel", "textnorm", "fluency"):
        all_keys = [k for k, _v in config.rule_items(kind)]
        miss = [k for k in all_keys if k not in hit_keys]
        print(f"{kind.upper():6} 命中 {len(all_keys) - len(miss)}/{len(all_keys)}"
              + (f"，未命中：{miss}" if miss else "，全部命中"))

    # ---- 新增模块验证：PDF / 自定义规则 / 自定义词库 ----
    print("\n" + "=" * 78)
    print("新增模块验证（PDF / 自定义规则 / 自定义词库）")
    print("=" * 78)
    by_name = {r.file_name: r for r in results}

    # 1) PDF 原生文本命中
    pdf_bad = by_name.get("sample_pdf_bad.pdf")
    if pdf_bad and pdf_bad.status != STATUS_UNREADABLE:
        print(f"  [PDF 文本检测] {pdf_bad.file_name} → {pdf_bad.status_text}，"
              f"命中 {len(pdf_bad.issues)} 条")
        loc_ok = any(i.location.startswith("第 ") and "页" in i.location
                     for i in pdf_bad.issues)
        print(f"    位置格式「第 N 页」：{'OK' if loc_ok else '缺失！'}")
    else:
        print("  [PDF 文本检测] sample_pdf_bad.pdf 缺失或判定为无法解析（异常）")

    # 2) 扫描图片型 PDF → 无法解析
    pdf_scan = by_name.get("sample_pdf_scan.pdf")
    if pdf_scan and pdf_scan.status == STATUS_UNREADABLE:
        print(f"  [扫描图片型 PDF] {pdf_scan.file_name} → 无法解析（隔离 OK）：{pdf_scan.error_message}")
    else:
        print("  [扫描图片型 PDF] sample_pdf_scan.pdf 未被判定为无法解析（异常）")

    # 3) 加密 PDF → 无法解析
    pdf_enc = by_name.get("sample_pdf_encrypted.pdf")
    if pdf_enc and pdf_enc.status == STATUS_UNREADABLE:
        print(f"  [加密 PDF] {pdf_enc.file_name} → 无法解析（隔离 OK）：{pdf_enc.error_message}")
    else:
        print("  [加密 PDF] sample_pdf_encrypted.pdf 未被判定为无法解析（异常）")

    # 4) 自定义规则命中（category=custom_rule）
    cust_hits = [i for r in results for i in r.issues if i.category == "custom_rule"]
    print(f"  [自定义规则] 命中 {len(cust_hits)} 条"
          + (f"，示例：{cust_hits[0].location} | {cust_hits[0].detail[:40]}"
             if cust_hits else "（未见命中！）"))
    # 期望样例命中：咱们 / 全角逗号后空格
    cust_patterns = {i.rule_title for i in cust_hits}
    for expect in ("禁用口语词『咱们』", "禁止使用全角逗号后接空格"):
        print(f"    · 期望规则「{expect}」：{'命中' if expect in cust_patterns else '未命中！'}")

    # 5) 自定义词库命中（category=wordbank）
    wb_hits = [i for r in results for i in r.issues if i.category == "wordbank"]
    print(f"  [自定义词库] 命中 {len(wb_hits)} 条"
          + (f"，示例：{wb_hits[0].location} | {wb_hits[0].detail[:40]}"
             if wb_hits else "（未见命中！）"))
    wb_kw = {i.rule_title for i in wb_hits}
    for expect in ("_tmp", "TODO"):
        print(f"    · 期望词条「{expect}」：{'命中' if expect in wb_kw else '未命中！'}")

    # ---- 报告生成 ----
    os.makedirs(OUT_DIR, exist_ok=True)
    out_path = os.path.join(OUT_DIR, "自检报告.docx")
    path = build_report(
        results, summary, out_path,
        rule_summary={
            "word": f"{config.enabled_count('word')}/{config.total_count('word')} 项",
            "excel": f"{config.enabled_count('excel')}/{config.total_count('excel')} 项",
            "textnorm": f"{config.enabled_count('textnorm')}/{config.total_count('textnorm')} 项",
            "pdf": f"{config.enabled_count('textnorm')}/{config.total_count('textnorm')} 项（同文字规范）",
        },
        operator="自动化自检", org="本地离线测试",
    )
    print("\n" + "=" * 78)
    print(f"报告已生成：{path}  ({os.path.getsize(path) / 1024:.1f} KB)")

    # 复核报告可被 python-docx 重新打开（验证文件结构有效）
    from docx import Document
    doc = Document(path)
    print(f"报告校验：段落 {len(doc.paragraphs)} 个，表格 {len(doc.tables)} 个 → 结构有效")
    print("=" * 78)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
