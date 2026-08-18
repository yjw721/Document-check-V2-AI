# -*- coding: utf-8 -*-
"""
测试样例生成脚本（本地运行，用于验证检测引擎命中情况）
=========================================================================
生成文件保存在 tests/samples/ ：
    sample_word_bad.docx    含大量 Word 低级错误
    sample_word_good.docx   规范文档（应零命中）
    sample_excel_bad.xlsx   含大量 Excel 低级错误
    sample_excel_good.xlsx  规范表格（应零命中）
    broken.docx             人为损坏文件（验证「无法解析」提示）
    legacy.doc              旧格式文件（验证格式不支持提示）

用法：python tests/make_samples.py
"""

from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment

SAMPLE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "samples")


def make_word_bad(path: str) -> None:
    """生成含各类低级错误的 Word 文档。"""
    doc = Document()

    doc.add_heading("项目实施方案", level=0)
    doc.add_heading("", level=1)                       # 空标题
    doc.add_heading("一、项目背景", level=1)

    # 全角空格 + 段首空格 + 段尾空格
    doc.add_paragraph("\u3000\u3000本项目\u3000旨在提升办公文档规范性,减少低级错误。  ")
    # 中英文标点混用 + 连续重复标点
    doc.add_paragraph("经研究决定,自2026年起全面推行文档核验制度。。请各部门认真落实!!")
    # 无效空白字符（零宽空格 + 不换行空格）
    doc.add_paragraph("质量控制\u200b要求：所有对外文件\u00a0必须完成核验后方可报送。")
    # 连续多个半角空格
    doc.add_paragraph("检查项目包括：格式规范    内容准确    数据一致")
    # 括号不配对
    doc.add_paragraph("详见附件（项目实施细则，由质量管理部负责解释。")

    # 连续空段落
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    doc.add_heading("二、工作要求", level=1)
    # 手工序号不连续（1 2 4 4）+ 格式混乱（混用「1.」与「1、」与「(1)」）
    doc.add_paragraph("1. 明确责任分工，落实到人。")
    doc.add_paragraph("2. 建立台账，逐项销号。")
    doc.add_paragraph("4. 定期开展抽查复核。")
    doc.add_paragraph("4. 结果纳入年度考核。")
    doc.add_paragraph("6、加强业务培训。")
    doc.add_paragraph("（1）每季度至少一次。")

    # 自动编号被普通段落打断
    p1 = doc.add_paragraph("第一阶段：制度建设", style="List Number")
    doc.add_paragraph("说明：本阶段为准备期，需完成制度起草与征求意见。")
    doc.add_paragraph("第二阶段：全面推行", style="List Number")

    # 段内大量手动换行
    p = doc.add_paragraph("联系方式：")
    for text in ["办公室：内网 8001", "质量管理部：内网 8002",
                 "信息中心：内网 8003", "值班电话：内网 8000"]:
        p.add_run().add_break()
        p.add_run(text)

    doc.add_heading("三、检查台账", level=1)
    # 表格：空白行 + 空白列 + 表头空单元格 + 行数多但无重复表头
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["序号", "部门", "", "检查结果", "备注"]   # 第 3 列表头为空
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i in range(1, 15):
        cells = table.add_row().cells
        if i == 7:
            continue          # 整行空白
        cells[0].text = str(i)
        cells[1].text = f"第 {i} 部门"
        cells[3].text = "合格" if i % 3 else "待整改"
        # 第 3 列与第 5 列始终为空 -> 空白列

    # 文字规范 / 表述问题检测样例（触发 7 类疑似规则）
    doc.add_heading("四、文字规范核查样例", level=1)
    # 触发 口语化 / 不规范用词：立马、搞定、事儿、甭、挺多
    doc.add_paragraph("我们立马搞定这个事儿，后续甭担心，相关事宜挺多要处理。")
    # 触发 重复冗余词句：亲眼目睹、凯旋归来、涉及到、目的是为了
    doc.add_paragraph("他亲眼目睹了凯旋归来的场面，本次工作涉及到多个部门，目的是为了实现目标。")
    # 触发 近义词误用：制定/制订、权利/权力、必须/必需
    doc.add_paragraph("公司制定了权力清单，必须明确各方权利与应尽义务。")
    # 触发 歧义 / 过长无标点长句（>80 字，无句末标点）+ 模糊词「相关」
    doc.add_paragraph(
        "本项目旨在通过对现有业务系统进行全面的升级与改造从而实现整体性能提升和稳定性增强"
        "并确保后续可维护性与安全性同时降低相关运维成本并提升用户满意度与系统可用性水平"
    )
    doc.add_paragraph("相关部门应在适当时机采取若干措施予以落实。")
    # 触发 错别字：安祥→安详、通谍→通牒
    doc.add_paragraph("部署工作已经安祥完成，本次通谍已下发，各项指标达到预期。")
    # 触发 非正式简称 / 自创缩写：发改委（禁用简称）、ZKB（未定义缩写，API/CRM 在白名单）
    doc.add_paragraph("经发改委审批后，采用ZKB结算机制，并与API、CRM系统完成对接。")
    # 触发 数量 / 单位表述不统一：万元 与 元 混用、5KM 非标准单位
    doc.add_paragraph("项目总投资 500万元，其中设备采购款 300元，线路总长度 5KM。")
    # 触发 英文拼写错误 / 英文语法错误 / 英文词汇不当（新增 3 类规则）
    doc.add_paragraph(
        "The goverment has recieved your proposal and will definately respond in time. "
        "He have went to the meeting, and I is not sure about the result. "
        "We don't have no time for a very unique solution, irregardless of the cost."
    )
    # 触发 中文语法错误：句式杂糅（根据…显示 / 通过…使）+ 关联词误配（只要…才）
    doc.add_paragraph("根据本次调研结果显示，通过优化流程使效率显著提升，只要认真整改，问题才会解决。")
    # 触发 中文词汇搭配不当：改善…水平、提高…力度、因为…的原因
    doc.add_paragraph("公司应改善服务水平并提高支持力度，因为时间紧迫的原因，需尽快落实。")
    # 触发 资产评估术语表述不规范：现金流折现法→现金流量折现法、评估基准日期→评估基准日、委托方→委托人
    doc.add_paragraph(
        "本次采用现金流折现法进行测算，评估基准日期为2025年12月31日，"
        "委托方要求按重置成本法另行复核一遍，收益现值法结果作为参考。"
    )
    # 触发 语句通顺度 6 类：逻辑断裂 / 成分残缺 / 语序混乱 / 重复赘述 / 关联词搭配 / 句式杂糅
    doc.add_paragraph(
        "因为时间紧迫。通过本次调研。而且我们要做好统筹工作不仅要注意细节。"
        "为了提升效率为了降低成本，必须优化流程。这项工作既重要而且紧迫。"
        "导致该项目延误的原因是由于管理混乱。"
    )

    # 文档尾部大量空段落
    for _ in range(4):
        doc.add_paragraph("")

    doc.save(path)


def make_word_good(path: str) -> None:
    """生成规范的 Word 文档（预期零命中或极少命中）。"""
    doc = Document()
    doc.add_heading("文档规范化操作指引", level=0)
    doc.add_heading("一、总体要求", level=1)
    doc.add_paragraph("文档排版应使用样式控制，避免手工插入空格与空行调整版式。")
    doc.add_paragraph("正文标点统一使用中文全角标点，数字与英文之间保持规范间距。")
    doc.add_heading("二、具体规范", level=1)
    doc.add_paragraph("一、标题层级应连续，不得出现空标题。")
    doc.add_paragraph("二、序号建议使用自动编号，避免人工维护出错。")
    doc.add_paragraph("三、表格应设置标题行跨页重复，保证可读性。")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["序号", "检查项", "标准"]):
        table.rows[0].cells[i].text = h
    for i in range(1, 4):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = f"检查项 {i}"
        cells[2].text = "符合规范"
    doc.save(path)


def make_excel_bad(path: str) -> None:
    """生成含各类低级错误的 Excel 工作簿。"""
    wb = Workbook()

    # ---- 表 1：数据表（多种问题）----
    ws = wb.active
    # 名称含全角空格 + 尾部空格（Excel 允许但属不规范；[ ] 等非法字符 openpyxl 禁止写入）
    ws.title = "数据表\u3000汇总 "
    ws.append(["序号", "部门", "", "金额", "金额", "完成率"])   # C 列表头空 + 金额重复
    rows = [
        [1, "办公室", None, "1200", 1200, 0.85],       # 数字文本化
        [2, "财务部 ", None, "3500", 3500, 0.92],      # 单元格尾部空格
        [3, "信息中心", None, 4800, 4800, 0.78],
        [None, None, None, None, None, None],          # 整行空白
        [4, "质量管理部", None, "文本值", 5200, 0.88],  # 类型混排
        [5, "人事部\u3000", None, "7,800", 7800, 0.95],  # 全角空格 + 千分位文本
    ]
    for r in rows:
        ws.append(r)

    # 公式错误
    ws["H1"] = "校验"
    ws["H2"] = "=1/0"
    ws["H3"] = "=NA()"
    ws["H4"] = "=UNKNOWNFUNC(1)"
    # openpyxl 不计算公式，手工写入错误值到另一列，模拟已缓存的错误结果
    ws["I1"] = "缓存结果"
    ws["I2"] = "#DIV/0!"
    ws["I3"] = "#N/A"
    ws["I4"] = "#NAME?"
    ws["I5"] = "#VALUE!"
    ws["I6"] = "#REF!"

    # 空内容合并单元格 + 数据区合并单元格
    ws.merge_cells("K2:L2")
    ws.merge_cells("A9:B9")
    ws["A9"] = "小计"

    # 隐藏行列
    ws.row_dimensions[6].hidden = True
    ws.column_dimensions["G"].hidden = True

    # 格式混乱：同列多种数字格式
    ws["F2"].number_format = "0.00%"
    ws["F3"].number_format = "0.0"
    ws["F4"].number_format = "General"
    ws["F5"].number_format = "#,##0.00"
    ws["F6"].alignment = Alignment(horizontal="right")

    # ---- 表 2：空白表 ----
    wb.create_sheet("Sheet2")

    # ---- 表 3：名称含首尾空格 ----
    wb.create_sheet(" 明细表 ")

    # ---- 表 4：整列空白 ----
    ws4 = wb.create_sheet("统计表")
    ws4.append(["项目", "空列", "数量"])
    for i in range(1, 6):
        ws4.append([f"项目{i}", None, i * 10])

    wb.save(path)


def make_excel_good(path: str) -> None:
    """生成规范的 Excel 工作簿。"""
    wb = Workbook()
    ws = wb.active
    ws.title = "部门完成情况"
    ws.append(["序号", "部门", "金额", "完成率"])
    data = [
        [1, "办公室", 1200, 0.85],
        [2, "财务部", 3500, 0.92],
        [3, "信息中心", 4800, 0.78],
        [4, "质量管理部", 5200, 0.88],
    ]
    for r in data:
        ws.append(r)
    for row in range(2, 6):
        ws.cell(row=row, column=3).number_format = "#,##0"
        ws.cell(row=row, column=4).number_format = "0.00%"
    wb.save(path)


def make_word_broken_image(path: str) -> None:
    """
    生成含损坏图片的 Word 文档：
    先正常插入一张图片，再把 ZIP 内的媒体文件替换为 0 字节，模拟图片加载异常。
    """
    import io
    import zipfile

    from docx.shared import Inches

    # 最小合法 PNG（1x1 像素，纯本地构造，无需任何外部素材）
    png_bytes = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
        "890000000a49444154789c63000100000500010d0a2db40000000049454e44ae426082"
    )

    doc = Document()
    doc.add_heading("含图片异常的文档", level=0)
    doc.add_paragraph("下方图片在保存后被人为破坏，用于验证「图片/对象加载异常」检测项。")
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(1.2))
    doc.add_paragraph("图 1  示例图片")

    # 全程内存操作，不产生任何临时文件
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    # 重打包：把 word/media/ 下的图片写成 0 字节
    with zipfile.ZipFile(buf) as src, zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.startswith("word/media/"):
                data = b""          # 制造 0 字节图片
            dst.writestr(item, data)


def make_broken(path: str) -> None:
    """生成人为损坏的 docx（ZIP 头正确但内容损坏）。"""
    with open(path, "wb") as fp:
        fp.write(b"PK\x03\x04" + os.urandom(220))


def make_legacy(path: str) -> None:
    """生成模拟旧格式 OLE 文件（.doc）。"""
    with open(path, "wb") as fp:
        fp.write(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 500)


# ---------------------------------------------------------------------------
# PDF 样例生成（本地离线、无额外依赖）
#   采用手工拼装的最小 PDF：文本以 UTF-16BE（带 BOM）十六进制串写入内容流，
#   不依赖 reportlab 等第三方绘图库；pypdf 的 extract_text() 可正确还原 Unicode，
#   满足「仅提取原生文本层、不 OCR」的检测约束。
# ---------------------------------------------------------------------------
def _utf16be_hex(text: str) -> str:
    """把文本编码为 UTF-16BE 码元序列的十六进制串（Identity-H 下 CID = 码元）。"""
    return "".join(f"{ord(c):04X}" for c in text)


def _build_pdf(pages_text: List[str], path: str) -> None:
    """
    生成多页 PDF（本地离线、无额外依赖）。

    采用 CID 字体（/Subtype /Type0 + /Encoding /Identity-H）+ ToUnicode CMap，
    使 pypdf 的 extract_text() 能正确还原中文 Unicode，满足「仅提取原生文本层、
    不 OCR」的检测约束。文本层以 UTF-16BE 码元十六进制串写入内容流。
    """
    n_pages = len(pages_text)

    # 收集全部去重字符，构造 ToUnicode CMap（每个 CID 直接映射回自身码点）
    seen: set = set()
    chars: List[str] = []
    for pg in pages_text:
        for ch in pg:
            if ch not in seen:
                seen.add(ch)
                chars.append(ch)
    bfchar = "\n".join(f"<{ord(c):04X}> <{ord(c):04X}>" for c in chars)
    tou_stream = (
        "/CIDInit /ProcSet findresource begin\n"
        "12 dict begin\n"
        "begincmap\n"
        "/CIDSystemInfo << /Registry (Adobe) /Ordering (Identity) /Supplement 0 >> def\n"
        "/CMapName /Adobe-Identity-UCS def\n"
        "/CMapType 2 def\n"
        "1 begincodespacerange\n"
        "<0000> <FFFF>\n"
        "endcodespacerange\n"
        f"{len(chars)} beginbfchar\n"
        f"{bfchar}\n"
        "endbfchar\n"
        "endcmap\n"
        "CMapName currentdict /CMap defineresource pop\n"
        "end\n"
        "end\n"
    )
    tou_bytes = tou_stream.encode("latin-1")

    # 对象编号：1 Catalog / 2 Pages / 3 Type0 字体 / 4 CIDFontType2 / 5 ToUnicode
    font_num = 3
    desc_num = 4
    tou_num = 5
    page_obj_nums: List[int] = []
    content_obj_nums: List[int] = []
    next_num = 6
    for _ in range(n_pages):
        page_obj_nums.append(next_num); next_num += 1
        content_obj_nums.append(next_num); next_num += 1
    total_objs = next_num - 1

    kids = " ".join(f"{n} 0 R" for n in page_obj_nums)

    objs: List[tuple] = []
    objs.append((1, "<< /Type /Catalog /Pages 2 0 R >>"))
    objs.append((2, f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>"))
    objs.append((font_num,
                 "<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light "
                 "/Encoding /Identity-H "
                 f"/DescendantFonts [{desc_num} 0 R] /ToUnicode {tou_num} 0 R >>"))
    objs.append((desc_num,
                 "<< /Type /Font /Subtype /CIDFontType2 /BaseFont /STSong-Light "
                 "/CIDSystemInfo << /Registry (Adobe) /Ordering (Identity) /Supplement 0 >> "
                 "/CIDToGIDMap /Identity >>"))
    objs.append((tou_num, f"<< /Length {len(tou_bytes)} >>\nstream\n{tou_stream}\nendstream"))

    for i in range(n_pages):
        raw_lines = pages_text[i].split("\n")
        y = 800
        parts: List[str] = []
        for line in raw_lines:
            if line == "":
                y -= 18
                continue
            parts.append(f"BT /F1 11 Tf 1 0 0 1 50 {y} Tm <{_utf16be_hex(line)}> Tj ET")
            y -= 18
        stream = "\n".join(parts)
        stream_bytes = stream.encode("latin-1")
        page_num = page_obj_nums[i]
        content_num = content_obj_nums[i]
        page_dict = (f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                     f"/Resources << /Font << /F1 {font_num} 0 R >> >> "
                     f"/Contents {content_num} 0 R >>")
        objs.append((page_num, page_dict))
        objs.append((content_num, f"<< /Length {len(stream_bytes)} >>\nstream\n{stream}\nendstream"))

    parts_out: List[str] = ["%PDF-1.4\n"]
    offsets: Dict[int, int] = {}
    pos = len(parts_out[0].encode("latin-1"))
    for num, payload in objs:
        obj_str = f"{num} 0 obj\n{payload}\nendobj\n"
        offsets[num] = pos
        parts_out.append(obj_str)
        pos += len(obj_str.encode("latin-1"))

    xref_offset = pos
    xref = f"xref\n0 {total_objs + 1}\n0000000000 65535 f \n"
    for num in range(1, total_objs + 1):
        xref += f"{offsets.get(num, 0):010d} 00000 n \n"
    parts_out.append(xref)
    trailer = (f"trailer\n<< /Size {total_objs + 1} /Root 1 0 R >>\n"
               f"startxref\n{xref_offset}\n%%EOF\n")
    parts_out.append(trailer)

    with open(path, "wb") as fp:
        fp.write("".join(parts_out).encode("latin-1"))


def make_pdf_bad(path: str) -> None:
    """生成含可提取文本层的 PDF：触发文字规范 + 自定义规则 + 自定义词库命中。"""
    pages = [
        # 第 1 页：模糊词 + 自定义规则「全角逗号后接空格」
        "我们决定， 重新规划本项目时间表。\n"
        "相关部门应在适当时机采取若干措施予以落实。",
        # 第 2 页：口语化 + 错别字 + 单位混用 + 自定义规则「禁用口语词『咱们』」
        "咱们立马搞定这个事儿，请各部门配合推进。\n"
        "项目总投资 500万元，其中设备采购 300元，线路总长 5KM。\n"
        "安祥地完成部署工作，各项指标达到预期。",
        # 第 3 页：自定义词库命中（_tmp / TODO 占位符）
        "本页含临时标记 _tmp 与待办 TODO，正式归档前请清理。",
    ]
    _build_pdf(pages, path)


def make_pdf_scan(path: str) -> None:
    """生成扫描图片型 PDF（内容流为空，无可提取文本层）→ 应判定为无法解析。"""
    # 每页仅保留空文本对象，模拟扫描件没有文字层
    _build_pdf(["\n\n\n", "\n\n\n"], path)


def make_pdf_encrypted(path: str) -> None:
    """生成加密 PDF（需密码）→ 应判定为无法解析。优先用 pypdf 加密，失败则退化为明文。"""
    tmp = path + ".clear.pdf"
    _build_pdf(["这是一份用于加密测试的普通文档，设置打开密码后应无法解析。"], tmp)
    try:
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(tmp)
        writer = PdfWriter()
        writer.append(reader)
        writer.encrypt("1234")
        with open(path, "wb") as fp:
            writer.write(fp)
        if os.path.exists(tmp):
            os.remove(tmp)
    except Exception as exc:  # noqa: BLE001
        if os.path.exists(tmp) and not os.path.exists(path):
            os.replace(tmp, path)
        print(f"  [提示] 加密 PDF 生成失败，已退化为明文样例：{exc}")


def main() -> None:
    os.makedirs(SAMPLE_DIR, exist_ok=True)
    tasks = [
        ("sample_word_bad.docx", make_word_bad),
        ("sample_word_good.docx", make_word_good),
        ("sample_word_image_bad.docx", make_word_broken_image),
        ("sample_excel_bad.xlsx", make_excel_bad),
        ("sample_excel_good.xlsx", make_excel_good),
        ("broken.docx", make_broken),
        ("legacy.doc", make_legacy),
        ("sample_pdf_bad.pdf", make_pdf_bad),
        ("sample_pdf_scan.pdf", make_pdf_scan),
        ("sample_pdf_encrypted.pdf", make_pdf_encrypted),
    ]
    for name, func in tasks:
        path = os.path.join(SAMPLE_DIR, name)
        func(path)
        print(f"[生成] {name}  ({os.path.getsize(path)} 字节)")
    print(f"\n样例目录：{SAMPLE_DIR}")


if __name__ == "__main__":
    main()
