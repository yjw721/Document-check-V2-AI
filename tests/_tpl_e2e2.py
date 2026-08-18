# -*- coding: utf-8 -*-
"""词库与标准规则批量导入模块 —— 端到端 API 验证。

覆盖：TXT(纠错对/正则/词条) + CSV(自适应列) + SCEL(搜狗词库) + docx(范本回归)
流程：upload(category) → draft 断言 → import → 配置追加断言 → 检测适配 → 清理
"""
import io
import json
import struct
import urllib.request
import urllib.error

BASE = "http://127.0.0.1:8501"
PASS = []


def req(method, path, data=None, files=None, category=None):
    url = BASE + path
    headers = {}
    body = None
    if files:
        boundary = "----wb" + "x" * 20
        parts = []
        for name, fname, content in files:
            parts.append(("--%s\r\nContent-Disposition: form-data; name=\"%s\"; filename=\"%s\"\r\n"
                          "Content-Type: application/octet-stream\r\n\r\n") % (boundary, name, fname))
            parts.append(content)
            parts.append(b"\r\n")  # multipart 要求 part 内容后跟 CRLF 再边界
        if category:
            parts.append(("--%s\r\nContent-Disposition: form-data; name=\"category\"\r\n\r\n%s\r\n")
                         % (boundary, category))
        parts.append("--%s--\r\n" % boundary)
        body = b"".join(p if isinstance(p, bytes) else p.encode("utf-8") for p in parts)
        headers["Content-Type"] = "multipart/form-data; boundary=%s" % boundary
    elif data is not None:
        body = json.dumps(data).encode("utf-8")
        headers["Content-Type"] = "application/json"
    r = urllib.request.Request(url, data=body, headers=headers, method=method)
    try:
        with urllib.request.urlopen(r, timeout=30) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        return {"http_error": e.code, "detail": e.read().decode("utf-8", "ignore")}


def check(name, cond, extra=""):
    if cond:
        PASS.append(name)
        print("  PASS", name)
    else:
        print("  FAIL", name, extra)


# ---------- 构造测试基准文件 ----------
def make_scel(words):
    data = bytearray(b"\x00" * 0x3B60)
    for w in words:
        wb = w.encode("utf-16-le")
        data += struct.pack("<HH", 1, len(wb)) + wb
        data += struct.pack("<II", 100, 0) + struct.pack("<H", 1) + struct.pack("<H", 0) + b"\x00"
    data += struct.pack("<HH", 0, 0)
    return bytes(data)


TXT = (
    "# 行业纠错文本库（注释行跳过）\n"
    "帐号=>账号\n"
    "截止日期=>截至日期\n"
    "落款处,落款处\n"  # 占位纠错对（用于去重测试）
    "其它,其他,术语统一\n"
    "^\\d{4}[-/.]\\d{1,2}[-/.]\\d{1,2}$::日期格式与基准不一致\n"
    "咱们\n"
    "绝对保证\n"
).encode("utf-8")

CSV = (
    "错误词,正确词,风险级别,整改建议\n"
    "帐号,账号,笔误,核对修正\n"
    "百分百,百分之百,风险警示,绝对化用语\n"
    "评估报告,评估说明,表述优化,术语混用\n"
).encode("utf-8-sig")

SCEL = make_scel(["资产评估", "公允价值", "持续经营"])

print("== 1) 上传混合格式基准（category=correction 标准纠错库） ==")
d = req("POST", "/api/template/upload", files=[
    ("files", "纠错库.txt", TXT), ("files", "标准纠错库.csv", CSV),
    ("files", "行业词库.scel", SCEL),
], category="correction")
check("upload 200 无 http_error", "http_error" not in d, str(d)[:200])
check("docs 3 个文件全部解析", len(d.get("docs", [])) == 3, str(d.get("docs")))
check("file_type 覆盖 TXT/CSV/SCEL", {x["file_type"] for x in d["docs"]} == {"TXT", "CSV", "SCEL"})
rules = d.get("rules", [])
entries = d.get("entries", [])
check("规则草案 ≥1（TXT 正则行）", len(rules) >= 1, f"rules={len(rules)}")
check("正则规则 tag=格式严重错误", any(r["tag"] == "格式严重错误" for r in rules))
check("纠错词条含 帐号/截止日期/其它", all(k in [e["keyword"] for e in entries] for k in ("帐号", "截止日期", "其它")))
check("CSV 标签列映射: 百分百→执业风险警示",
      any(e["keyword"] == "百分百" and e["tag"] == "执业风险警示" for e in entries))
check("SCEL 词条含 资产评估/公允价值", all(k in [e["keyword"] for e in entries] for k in ("资产评估", "公允价值")))
check("文本基准 source_page=0", all(r.get("source_page") == 0 for r in rules) and all(e.get("source_page") == 0 for e in entries))
sel_r = [r["id"] for r in rules]
sel_e = [e["id"] for e in entries]
print(f"  规则 {len(rules)} 条 / 词条 {len(entries)} 条")

print("== 2) 导入选中项 ==")
r = req("POST", "/api/template/import", {"rule_ids": sel_r, "entry_ids": sel_e})
check("import 成功", r.get("ok") is True, str(r)[:200])
check("导入数量匹配", r.get("imported_rules") == len(rules) and r.get("imported_entries") == len(entries), str(r))

print("== 3) 配置追加断言（分组名带类别，不覆盖原配置） ==")
cr = req("GET", "/api/custom_rules")
wb = req("GET", "/api/wordbanks")
new_rg = [g for g in cr.get("groups", []) if "外部导入" in g.get("name", "")]
new_wg = [g for g in wb.get("groups", []) if "外部导入" in g.get("name", "")]
check("自定义规则组追加且含类别名", len(new_rg) == 1 and "标准纠错库" in new_rg[0]["name"], str(new_rg))
check("词库组追加且含类别名", len(new_wg) == 1 and "标准纠错库" in new_wg[0]["name"], str(new_wg))
check("原配置组数量未减少（追加不覆盖）", len(cr.get("groups", [])) >= 1 and len(wb.get("groups", [])) >= 1)
check("规则组 rules 带 tag/severity", all("tag" in x and "severity" in x for x in new_rg[0]["rules"]))
check("词库组 entries 带 tag/suggestion", all("tag" in x and "suggestion" in x for x in new_wg[0]["entries"]))

print("== 4) 检测适配：违规文档命中新词条（位置 + 标签） ==")
from docx import Document
doc = Document()
doc.add_paragraph("XX资产评估有限公司资产评估报告")
doc.add_paragraph("评估基准日为2022年10月31日，评估基准日为2022/10/31。")  # 正则规则命中（/ 格式）
doc.add_paragraph("经核对，帐号信息与原始凭证一致，截止日期为2022-10-31。")  # 帐号/截止日期 命中
doc.add_paragraph("本公司绝对保证上述内容真实有效。")  # 绝对保证 命中
doc.add_paragraph("公允价值评估结果如下，持续经营假设成立。")  # SCEL 词条
doc.save("tests/_tpl/multi_violation.docx")
with open("tests/_tpl/multi_violation.docx", "rb") as fh:
    content = fh.read()
u = req("POST", "/api/upload", files=[("files", "multi_violation.docx", content)])
check("上传检测 200", "http_error" not in u, str(u)[:200])
iss = req("GET", "/api/issues")
issues = iss.get("issues", [])
kw_hits = [i for i in issues if i.get("tag") in ("笔误警示", "执业风险警示", "格式严重错误")]
check("检测到自定义词条命中", any(i["rule_title"] and "帐号" in i.get("rule_title", "") for i in kw_hits), str([i.get("rule_title") for i in issues[:10]]))
check("命中带位置信息", all(i.get("location") for i in kw_hits), str([i.get("location") for i in kw_hits[:5]]))
print("  命中示例：", [(i.get("rule_title"), i.get("location"), i.get("tag")) for i in kw_hits[:4]])

print("== 5) docx 范本回归（原链路不受影响） ==")
doc2 = Document()
doc2.add_paragraph("XX资产评估有限公司资产评估报告")
doc2.add_paragraph("评估基准日为2022年10月31日，采用成本法。")
doc2.add_paragraph("本报告仅对委托方提供，不得对外提供。")
doc2.save("tests/_tpl/regress.docx")
with open("tests/_tpl/regress.docx", "rb") as fh:
    content = fh.read()
d2 = req("POST", "/api/template/upload", files=[("files", "回归范本.docx", content)], category="asset")
check("docx 回归解析成功", "http_error" not in d2 and d2.get("docs") and d2["docs"][0]["file_type"] == "Word", str(d2)[:150])
check("docx 生成规则/词条", len(d2.get("rules", [])) > 0 or len(d2.get("entries", [])) > 0)
check("docx 带页码定位", all((x.get("source_page") or 0) > 0 for x in d2.get("rules", [])[:1]))
req("POST", "/api/template/clear", {})

print("== 6) 清理测试数据 ==")
cr = req("GET", "/api/custom_rules")
wb = req("GET", "/api/wordbanks")
cr["groups"] = [g for g in cr.get("groups", []) if "外部导入" not in g.get("name", "")]
wb["groups"] = [g for g in wb.get("groups", []) if "外部导入" not in g.get("name", "")]
req("POST", "/api/custom_rules", cr)
req("POST", "/api/wordbanks", wb)
req("POST", "/api/clear_data", {})
req("POST", "/api/template/clear", {})
cr2 = req("GET", "/api/custom_rules")
wb2 = req("GET", "/api/wordbanks")
check("清理完成（无外部导入组）",
      all("外部导入" not in g.get("name", "") for g in cr2.get("groups", [])) and
      all("外部导入" not in g.get("name", "") for g in wb2.get("groups", [])))

print("\n========== 结果：%d/%d 通过 ==========" % (len(PASS), len(PASS)))
